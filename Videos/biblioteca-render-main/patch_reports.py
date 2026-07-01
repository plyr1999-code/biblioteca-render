import os

base_path = r'c:\Users\user\biblioteca\biblioteca_python\routes'

configs = {
    'libros.py': {
        'bp': 'libros_bp',
        'model': 'LibrosModel',
        'get_method': 'get_libros',
        'title': 'Listado de Libros',
        'cols': [('N°', 10), ('Título', 55), ('Autor', 40), ('Materia', 40), ('Cant.', 15), ('Estado', 20)],
        'keys': ['titulo', 'autor', 'materia', 'cantidad', 'estado']
    },
    'materia.py': {
        'bp': 'materia_bp',
        'model': 'MateriaModel',
        'get_method': 'get_materia',
        'title': 'Listado de Materias',
        'cols': [('N°', 15), ('Materia', 130), ('Estado', 30)],
        'keys': ['materia', 'estado']
    },
    'editorial.py': {
        'bp': 'editorial_bp',
        'model': 'EditorialModel',
        'get_method': 'get_editorial',
        'title': 'Listado de Editoriales',
        'cols': [('N°', 15), ('Editorial', 130), ('Estado', 30)],
        'keys': ['editorial', 'estado']
    },
    'estudiantes.py': {
        'bp': 'estudiantes_bp',
        'model': 'EstudiantesModel',
        'get_method': 'get_estudiante',
        'title': 'Listado de Estudiantes',
        'cols': [('N°', 10), ('Código', 25), ('DNI', 25), ('Nombre', 50), ('Carrera', 45), ('Teléfono', 25)],
        'keys': ['codigo', 'dni', 'nombre', 'carrera', 'telefono']
    },
    'usuarios.py': {
        'bp': 'usuarios_bp',
        'model': 'UsuariosModel',
        'get_method': 'get_usuario',
        'title': 'Listado de Usuarios',
        'cols': [('N°', 10), ('Usuario', 30), ('Nombre', 50), ('Correo', 60), ('Estado', 30)],
        'keys': ['usuario', 'nombre', 'correo', 'estado']
    }
}

TEMPLATE = """

@{bp}.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir: return redir
    from fpdf import FPDF
    from flask import Response
    m = {model}()
    registros = m.{get_method}()
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('{title}')
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.cell(195, 8, '{title}', 0, 1, 'C')
    pdf_doc.ln(3)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(26, 58, 92)
    pdf_doc.set_text_color(255, 255, 255)
{pdf_headers}
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(registros, 1):
{pdf_rows}
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={{'Content-Disposition': 'inline; filename=reporte.pdf'}})


@{bp}.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response
    m = {model}()
    registros = m.{get_method}()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "{title}"
    header_fill = PatternFill("solid", fgColor="1a3a5c")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers_arr = [{excel_headers}]
    for col, h in enumerate(headers_arr, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for i, row in enumerate(registros, 1):
        ws.append([i, {excel_rows}])
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={{'Content-Disposition': f'attachment; filename=reporte_{{today}}.xlsx'}}
    )


@{bp}.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import Response
    m = {model}()
    registros = m.{get_method}()
    doc = Document()
    title = doc.add_heading('{title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Generado el {{date.today().isoformat()}}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    headers_arr = [{excel_headers}]
    table = doc.add_table(rows=1, cols=len(headers_arr))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers_arr):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs: run.bold = True
    for i, row in enumerate(registros, 1):
        cells = table.add_row().cells
        vals = [str(i), {word_rows}]
        for j, v in enumerate(vals):
            cells[j].text = v
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={{'Content-Disposition': f'attachment; filename=reporte_{{today}}.docx'}}
    )
"""

for filename, cfg in configs.items():
    filepath = os.path.join(base_path, filename)
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    if "@%s.route('/pdf')" % cfg['bp'] in content:
        print(f"Skipping {filename}, already has reports.")
        continue

    # Build PDF headers and rows
    pdf_headers = []
    pdf_rows = []
    excel_headers = []
    excel_rows = []
    word_rows = []

    for i, (col_name, width) in enumerate(cfg['cols']):
        excel_headers.append(repr(col_name))
        is_last = (i == len(cfg['cols']) - 1)
        term = '1' if is_last else '0'
        pdf_headers.append(f"    pdf_doc.cell({width}, 6, '{col_name}', 1, {term}, 'C', True)")
        
        if i == 0:
            pdf_rows.append(f"        pdf_doc.cell({width}, 6, str(i), 1, {term}, 'C')")
        else:
            k = cfg['keys'][i-1]
            if k == 'estado':
                pdf_rows.append(f"        pdf_doc.cell({width}, 6, 'Activo' if row.get('{k}') == 1 else 'Inactivo', 1, {term}, 'C')")
                excel_rows.append(f"'Activo' if row.get('{k}') == 1 else 'Inactivo'")
                word_rows.append(f"'Activo' if row.get('{k}') == 1 else 'Inactivo'")
            else:
                lalign = "'C'" if width < 30 else "'L'"
                pdf_rows.append(f"        pdf_doc.cell({width}, 6, str(row.get('{k}', '')), 1, {term}, {lalign})")
                excel_rows.append(f"row.get('{k}', '')")
                word_rows.append(f"str(row.get('{k}', ''))")
                
    filled = TEMPLATE.format(
        bp=cfg['bp'],
        model=cfg['model'],
        get_method=cfg['get_method'],
        title=cfg['title'],
        pdf_headers="\n".join(pdf_headers),
        pdf_rows="\n".join(pdf_rows),
        excel_headers=", ".join(excel_headers),
        excel_rows=", ".join(excel_rows),
        word_rows=", ".join(word_rows)
    )

    with open(filepath, 'a', encoding='utf-8') as f:
        f.write(filled)

    print(f"Patched {filename} successfully.")
