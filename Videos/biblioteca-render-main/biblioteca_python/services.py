from logger import get_logger

logger = get_logger('app')

class BaseService:
    """Servicio base para lógica de negocio"""

    def __init__(self, model):
        self.model = model

    def paginate(self, items, page=1, per_page=10):
        """Paginar resultados"""
        total = len(items) if isinstance(items, list) else 0
        start = (page - 1) * per_page
        end = start + per_page

        paginated = items[start:end] if isinstance(items, list) else []
        total_pages = (total + per_page - 1) // per_page

        return {
            'items': paginated,
            'page': page,
            'per_page': per_page,
            'total': total,
            'total_pages': total_pages,
            'has_next': page < total_pages,
            'has_prev': page > 1
        }

    def search(self, items, query, fields):
        """Buscar en items"""
        if not query:
            return items

        query = query.lower()
        results = []

        for item in items:
            for field in fields:
                if field in item:
                    value = str(item[field]).lower()
                    if query in value:
                        results.append(item)
                        break

        return results

    def filter_by_status(self, items, status=1):
        """Filtrar por estado"""
        if status is None:
            return items
        return [item for item in items if item.get('estado') == status]

    def handle_error(self, error, operation="operation"):
        """Manejo centralizado de errores"""
        logger.error(f"Error in {operation}: {error}")
        return {
            'success': False,
            'error': str(error),
            'code': 500
        }

class UsuariosService(BaseService):
    """Servicio de usuarios"""

    def listar_con_paginacion(self, page=1, per_page=10, search=None, status=None):
        try:
            users = self.model.get_usuarios()

            if status is not None:
                users = self.filter_by_status(users, status)

            if search:
                users = self.search(users, search, ['usuario', 'nombre'])

            return self.paginate(users, page, per_page)
        except Exception as e:
            return self.handle_error(e, "list_users_pagination")

    def crear_usuario(self, usuario, nombre, clave):
        try:
            result = self.model.registrar_usuario(usuario, nombre, clave)
            return {
                'success': result == 'ok',
                'message': 'Usuario creado' if result == 'ok' else 'Error al crear usuario',
                'code': 201 if result == 'ok' else 400
            }
        except Exception as e:
            return self.handle_error(e, "create_user")

    def actualizar_usuario(self, id_, usuario, nombre):
        try:
            result = self.model.modificar_usuario(usuario, nombre, id_)
            return {
                'success': result == 'modificado',
                'message': 'Usuario actualizado' if result == 'modificado' else 'Error',
                'code': 200 if result == 'modificado' else 400
            }
        except Exception as e:
            return self.handle_error(e, "update_user")

    def obtener_usuario(self, id_):
        try:
            user = self.model.editar_user(id_)
            return {
                'success': user is not None,
                'user': user if user else None,
                'code': 200 if user else 404
            }
        except Exception as e:
            return self.handle_error(e, "get_user")

    def desactivar_usuario(self, id_):
        try:
            result = self.model.accion_user(0, id_)
            return {
                'success': result == 1,
                'message': 'Usuario desactivado' if result == 1 else 'Error',
                'code': 200 if result == 1 else 400
            }
        except Exception as e:
            return self.handle_error(e, "deactivate_user")

class LibrosService(BaseService):
    """Servicio de libros"""

    def listar_con_paginacion(self, page=1, per_page=10, search=None, status=None):
        try:
            libros = self.model.get_libros()

            if status is not None:
                libros = self.filter_by_status(libros, status)

            if search:
                libros = self.search(libros, search, ['titulo', 'isbn'])

            return self.paginate(libros, page, per_page)
        except Exception as e:
            return self.handle_error(e, "list_books_pagination")



class EstudiantesService(BaseService):
    """Servicio de estudiantes"""

    def listar_con_paginacion(self, page=1, per_page=10, search=None, status=None):
        try:
            estudiantes = self.model.get_estudiantes()

            if status is not None:
                estudiantes = self.filter_by_status(estudiantes, status)

            if search:
                estudiantes = self.search(estudiantes, search, ['nombre', 'carnet'])

            return self.paginate(estudiantes, page, per_page)
        except Exception as e:
            return self.handle_error(e, "list_students_pagination")


class ReportService:
    """Servicio para generar reportes en diferentes formatos - Premium"""

    def __init__(self):
        self.logger = get_logger('app')
        self.app_name = "Sistema de Gestión de Biblioteca"
        self.app_version = "2.1.0"

    def generate_pdf(self, titulo, data, columns, filename="reporte"):
        """Generar PDF profesional y pulido con datos tabulares"""
        try:
            from fpdf import FPDF
            from datetime import datetime
            import os

            # Asegurar que tiene extensión
            if not filename.endswith('.pdf'):
                filename = f"{filename}.pdf"

            pdf = FPDF('P', 'mm', 'letter')
            pdf.add_page()
            pdf.set_margins(12, 12, 12)
            pdf.set_title(titulo)
            pdf.set_author(self.app_name)

            # Encabezado profesional
            pdf.set_font('Arial', 'B', 16)
            pdf.set_text_color(25, 50, 100)
            pdf.cell(0, 12, titulo, 0, 1, 'C')

            # Línea separadora
            pdf.set_draw_color(70, 130, 180)
            pdf.line(12, pdf.get_y(), 198, pdf.get_y())
            pdf.ln(3)

            # Información del reporte
            pdf.set_font('Arial', 'I', 9)
            pdf.set_text_color(100, 100, 100)
            fecha_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            pdf.cell(0, 4, f"Generado: {fecha_str}", 0, 1, 'L')
            pdf.cell(0, 4, f"Versión: {self.app_version}", 0, 1, 'L')
            pdf.ln(2)

            # Tabla con diseño profesional
            pdf.set_font('Arial', 'B', 10)
            pdf.set_text_color(255, 255, 255)
            pdf.set_fill_color(25, 50, 100)
            col_width = 186 / len(columns)

            # Encabezados con fondo
            for col in columns:
                col_name = str(col)[:20]
                pdf.cell(col_width, 8, col_name, 1, 0, 'C', fill=True)
            pdf.ln(8)

            # Datos con alternancia de colores
            pdf.set_font('Arial', '', 9)
            pdf.set_text_color(0, 0, 0)
            alternate_color = False

            for row_idx, row in enumerate(data[:100]):  # Máx 100 filas para mejor rendimiento
                if alternate_color:
                    pdf.set_fill_color(240, 245, 250)
                else:
                    pdf.set_fill_color(255, 255, 255)

                for col in columns:
                    value = str(row.get(col, ''))[:25]
                    pdf.cell(col_width, 7, value, 1, 0, 'L', fill=True)
                pdf.ln(7)
                alternate_color = not alternate_color

            # Pie de página
            pdf.ln(5)
            pdf.set_font('Arial', 'I', 8)
            pdf.set_text_color(150, 150, 150)
            pdf.cell(0, 4, f"Total de registros: {len(data)}", 0, 1, 'C')

            pdf.output(filename)
            self.logger.info(f"PDF generado exitosamente: {filename}")
            return True
        except Exception as e:
            self.logger.error(f"Error generando PDF: {e}")
            return False

    def generate_excel(self, titulo, data, columns, filename="reporte.xlsx"):
        """Generar Excel profesional y pulido con datos tabulares"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from datetime import datetime

            if not filename.endswith('.xlsx'):
                filename = f"{filename}.xlsx"

            wb = Workbook()
            ws = wb.active
            ws.title = "Reporte"

            # Título
            ws['A1'] = titulo
            ws['A1'].font = Font(bold=True, size=16, color="193264")
            ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
            ws.merge_cells('A1:Z1')
            ws.row_dimensions[1].height = 25

            # Fecha de generación
            fecha_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            ws['A2'] = f"Generado: {fecha_str} | Versión: {self.app_version}"
            ws['A2'].font = Font(italic=True, size=9, color="666666")
            ws['A2'].alignment = Alignment(horizontal="left")
            ws.merge_cells('A2:Z2')
            ws.row_dimensions[2].height = 15

            ws.row_dimensions[3].height = 5

            # Estilos para encabezados
            header_fill = PatternFill(start_color="1932C0", end_color="1932C0", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Encabezados
            for col_num, col in enumerate(columns, 1):
                cell = ws.cell(row=4, column=col_num)
                cell.value = col
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = thin_border

            ws.row_dimensions[4].height = 20

            # Datos con formato y estilos
            data_fill = PatternFill(start_color="F0F5FA", end_color="F0F5FA", fill_type="solid")
            data_font = Font(size=10)
            data_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

            for row_num, row_data in enumerate(data[:1000], 5):
                for col_num, col in enumerate(columns, 1):
                    cell = ws.cell(row=row_num, column=col_num)
                    cell.value = row_data.get(col, '')
                    cell.font = data_font
                    if row_num % 2 == 0:
                        cell.fill = data_fill
                    cell.alignment = data_alignment
                    cell.border = thin_border

                ws.row_dimensions[row_num].height = 18

            # Ajustar ancho de columnas automáticamente
            for col_num, col in enumerate(columns, 1):
                max_length = len(str(col)) + 2
                for row in ws.iter_rows(min_col=col_num, max_col=col_num, min_row=5):
                    try:
                        if len(str(row[0].value)) > max_length:
                            max_length = len(str(row[0].value)) + 2
                    except:
                        pass
                ws.column_dimensions[chr(64 + col_num)].width = min(max_length, 50)

            # Agregar información al final
            last_row = len(data) + 6
            ws[f'A{last_row}'] = f"Total de registros: {len(data)}"
            ws[f'A{last_row}'].font = Font(italic=True, size=9, color="666666")

            wb.save(filename)
            self.logger.info(f"Excel generado exitosamente: {filename}")
            return True
        except Exception as e:
            self.logger.error(f"Error generando Excel: {e}")
            return False

    def generate_word(self, titulo, data, columns, filename="reporte.docx"):
        """Generar Word profesional y pulido con datos tabulares"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from datetime import datetime

            if not filename.endswith('.docx'):
                filename = f"{filename}.docx"

            doc = Document()

            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Encabezado del documento
            title = doc.add_heading(titulo, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_format = title.paragraph_format
            title_format.space_before = Pt(12)
            title_format.space_after = Pt(6)
            for run in title.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(25, 50, 100)
                run.font.bold = True

            # Información del reporte
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fecha_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            info_run = info_para.add_run(f"Generado: {fecha_str}\nVersión: {self.app_version}")
            info_run.font.size = Pt(9)
            info_run.font.italic = True
            info_run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()  # Espaciador

            # Tabla de datos
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = 'Light Grid Accent 1'

            # Encabezados con formato
            header_cells = table.rows[0].cells
            for i, col in enumerate(columns):
                header_cells[i].text = str(col)
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)

                # Fondo azul para encabezados
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '193264')
                header_cells[i]._element.get_or_add_tcPr().append(shading_elm)

            # Datos con formato
            for row_data in data[:500]:
                row_cells = table.add_row().cells
                for i, col in enumerate(columns):
                    row_cells[i].text = str(row_data.get(col, ''))
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            # Pie de página
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(f"Total de registros: {len(data)}")
            footer_run.font.size = Pt(9)
            footer_run.font.italic = True
            footer_run.font.color.rgb = RGBColor(150, 150, 150)

            doc.save(filename)
            self.logger.info(f"Word generado exitosamente: {filename}")
            return True
        except Exception as e:
            self.logger.error(f"Error generando Word: {e}")
            return False

