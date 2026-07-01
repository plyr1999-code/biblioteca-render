from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template
from models.estudiantes_model import EstudiantesModel
from helpers import str_clean
from security import permission_required, SecurityUtils

estudiantes_bp = Blueprint('estudiantes', __name__)


def login_required():
    if not session.get('activo'):
        return redirect(url_for('index'))
    return None


@estudiantes_bp.route('/')
def index():
    redir = login_required()
    if redir:
        return redir
    id_user = session['id_usuario']
    m = EstudiantesModel()
    perm = m.verificar_permisos(id_user, "Estudiantes")
    if not perm and id_user != 1:
        return render_template('permisos.html')
    return render_template('estudiantes/index.html')


@estudiantes_bp.route('/listar')
@permission_required('Estudiantes')
def listar():
    redir = login_required()
    if redir:
        return redir
    estado = request.args.get('estado')
    m = EstudiantesModel()
    data = m.get_estudiantes()
    if estado is not None and estado != '':
        data = [r for r in data if str(r.get('estado', '')) == str(estado)]
    es_admin = session.get('id_usuario') == 1
    for row in data:
        row['estadoLabel'] = 'Activo' if row['estado'] == 1 else 'Inactivo'
        if not es_admin:
            row['acciones'] = 'solo_lectura'
        else:
            row['acciones'] = 'editar,eliminar' if row['estado'] == 1 else 'reingresar'
    return jsonify(data)


@estudiantes_bp.route('/registrar', methods=['POST'])
@permission_required('Estudiantes')
def registrar():
    redir = login_required()
    if redir:
        return redir
    codigo = str_clean(request.form.get('codigo', ''))
    dni = str_clean(request.form.get('dni', ''))
    nombre = str_clean(request.form.get('nombre', ''))
    carrera = str_clean(request.form.get('carrera', ''))
    direccion = str_clean(request.form.get('direccion', ''))
    telefono = str_clean(request.form.get('telefono', ''))
    id_ = str_clean(request.form.get('id', ''))

    if not codigo or not nombre:
        return jsonify({'msg': 'Los campos codigo y nombre son requeridos', 'icono': 'warning'})
    if not SecurityUtils.validate_student_code(codigo):
        return jsonify({'msg': 'El codigo debe tener entre 2 y 20 caracteres alfanumericos', 'icono': 'warning'})
    if dni and not SecurityUtils.validate_dni(dni):
        return jsonify({'msg': 'El DNI debe tener entre 6 y 20 caracteres validos', 'icono': 'warning'})
    if not SecurityUtils.validate_person_name(nombre, max_len=100):
        return jsonify({'msg': 'El nombre debe tener solo letras y espacios (3-100 caracteres)', 'icono': 'warning'})
    if carrera and not SecurityUtils.validate_generic_text(carrera, min_len=2, max_len=80):
        return jsonify({'msg': 'La carrera contiene caracteres no permitidos o excede el tamano permitido', 'icono': 'warning'})
    if direccion and not SecurityUtils.validate_generic_text(direccion, min_len=3, max_len=150):
        return jsonify({'msg': 'La direccion contiene caracteres no permitidos o excede el tamano permitido', 'icono': 'warning'})
    if telefono and not SecurityUtils.validate_phone(telefono):
        return jsonify({'msg': 'Ingresa un telefono valido', 'icono': 'warning'})

    m = EstudiantesModel()
    if id_ == '':
        data = m.insertar_estudiante(codigo, dni, nombre, carrera, direccion, telefono)
        msgs = {'ok': ('Estudiante registrado', 'success'),
                'existe': ('El codigo ya existe', 'warning'),
                'error': ('Error al registrar', 'error')}
        msg, icono = msgs.get(data, ('Error', 'error'))
        return jsonify({'msg': msg, 'icono': icono})
    else:
        data = m.actualizar_estudiante(codigo, dni, nombre, carrera, direccion, telefono, id_)
        if data == 'modificado':
            return jsonify({'msg': 'Estudiante modificado', 'icono': 'success'})
        return jsonify({'msg': 'Error al modificar', 'icono': 'error'})


@estudiantes_bp.route('/editar/<int:id_>')
@permission_required('Estudiantes')
def editar(id_):
    redir = login_required()
    if redir:
        return redir
    return jsonify(EstudiantesModel().edit_estudiante(id_))


@estudiantes_bp.route('/eliminar/<int:id_>')
@permission_required('Estudiantes')
def eliminar(id_):
    redir = login_required()
    if redir:
        return redir
    data = EstudiantesModel().estado_estudiante(0, id_)
    if data == 1:
        return jsonify({'msg': 'Estudiante dado de baja', 'icono': 'success'})
    return jsonify({'msg': 'Error al eliminar', 'icono': 'error'})


@estudiantes_bp.route('/reingresar/<int:id_>')
@permission_required('Estudiantes')
def reingresar(id_):
    redir = login_required()
    if redir:
        return redir
    data = EstudiantesModel().estado_estudiante(1, id_)
    if data == 1:
        return jsonify({'msg': 'Estudiante restaurado', 'icono': 'success'})
    return jsonify({'msg': 'Error al restaurar', 'icono': 'error'})


@estudiantes_bp.route('/buscarEstudiante')
@permission_required('Estudiantes')
def buscar_estudiante():
    redir = login_required()
    if redir:
        return redir
    valor = request.args.get('es', '')
    return jsonify(EstudiantesModel().buscar_estudiante(valor))


@estudiantes_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir: return redir
    from fpdf import FPDF
    from flask import Response
    estado = request.args.get('estado')
    m = EstudiantesModel()
    registros = m.get_estudiantes()
    if estado is not None and estado != '':
        registros = [r for r in registros if str(r.get('estado', '')) == str(estado)]
    filtro_lbl = ('Activos' if str(estado) == '1' else 'Inactivos') if (estado is not None and estado != '') else 'Todos'
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Listado de Estudiantes')
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.cell(195, 8, 'Listado de Estudiantes', 0, 1, 'C')
    pdf_doc.set_font('Arial', 'I', 9)
    pdf_doc.set_text_color(100, 100, 100)
    pdf_doc.cell(195, 5, f'Estado: {filtro_lbl}', 0, 1, 'C')
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.ln(3)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(26, 58, 92)
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(10, 6, 'N°', 1, 0, 'C', True)
    pdf_doc.cell(25, 6, 'Código', 1, 0, 'C', True)
    pdf_doc.cell(25, 6, 'DNI', 1, 0, 'C', True)
    pdf_doc.cell(50, 6, 'Nombre', 1, 0, 'C', True)
    pdf_doc.cell(45, 6, 'Carrera', 1, 0, 'C', True)
    pdf_doc.cell(25, 6, 'Teléfono', 1, 1, 'C', True)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(registros, 1):
        pdf_doc.cell(10, 6, str(i), 1, 0, 'C')
        pdf_doc.cell(25, 6, str(row.get('codigo', '')), 1, 0, 'C')
        pdf_doc.cell(25, 6, str(row.get('dni', '')), 1, 0, 'C')
        pdf_doc.cell(50, 6, str(row.get('nombre', '')), 1, 0, 'L')
        pdf_doc.cell(45, 6, str(row.get('carrera', '')), 1, 0, 'L')
        pdf_doc.cell(25, 6, str(row.get('telefono', '')), 1, 1, 'C')
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=reporte.pdf'})


@estudiantes_bp.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response
    estado = request.args.get('estado')
    m = EstudiantesModel()
    registros = m.get_estudiantes()
    if estado is not None and estado != '':
        registros = [r for r in registros if str(r.get('estado', '')) == str(estado)]
    filtro_lbl = ('Activos' if str(estado) == '1' else 'Inactivos') if (estado is not None and estado != '') else 'Todos'
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Listado de Estudiantes"
    ws.merge_cells('A1:F1')
    sub_cell = ws['A1']
    sub_cell.value = f'Estado: {filtro_lbl}'
    sub_cell.font = Font(italic=True, color='FF555555', size=10)
    sub_cell.alignment = Alignment(horizontal='center')
    header_fill = PatternFill("solid", fgColor="1a3a5c")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers_arr = ['N°', 'Código', 'DNI', 'Nombre', 'Carrera', 'Teléfono']
    for col, h in enumerate(headers_arr, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for i, row in enumerate(registros, 1):
        ws.append([i, row.get('codigo', ''), row.get('dni', ''), row.get('nombre', ''), row.get('carrera', ''), row.get('telefono', '')])
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.xlsx'}
    )


@estudiantes_bp.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import Response
    estado = request.args.get('estado')
    m = EstudiantesModel()
    registros = m.get_estudiantes()
    if estado is not None and estado != '':
        registros = [r for r in registros if str(r.get('estado', '')) == str(estado)]
    filtro_lbl = ('Activos' if str(estado) == '1' else 'Inactivos') if (estado is not None and estado != '') else 'Todos'
    doc = Document()
    title = doc.add_heading('Listado de Estudiantes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Generado el {date.today().isoformat()} | Estado: {filtro_lbl}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    headers_arr = ['N°', 'Código', 'DNI', 'Nombre', 'Carrera', 'Teléfono']
    table = doc.add_table(rows=1, cols=len(headers_arr))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers_arr):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for i, row in enumerate(registros, 1):
        cells = table.add_row().cells
        vals = [str(i), str(row.get('codigo', '')), str(row.get('dni', '')), str(row.get('nombre', '')), str(row.get('carrera', '')), str(row.get('telefono', ''))]
        for j, v in enumerate(vals):
            cells[j].text = v
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.docx'}
    )
