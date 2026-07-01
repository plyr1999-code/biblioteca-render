import os
from datetime import datetime
from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template, current_app
from models.autor_model import AutorModel
from helpers import str_clean
from security import permission_required, SecurityUtils

autor_bp = Blueprint('autor', __name__)


def login_required():
    if not session.get('activo'):
        return redirect(url_for('index'))
    return None


@autor_bp.route('/')
def index():
    redir = login_required()
    if redir:
        return redir
    id_user = session['id_usuario']
    m = AutorModel()
    perm = m.verificar_permisos(id_user, "Autor")
    if not perm and id_user != 1:
        return render_template('permisos.html')
    return render_template('autor/index.html')


@autor_bp.route('/listar')
@permission_required('Autor')
def listar():
    redir = login_required()
    if redir:
        return redir
    m = AutorModel()
    data = m.get_autor()
    es_admin = session.get('id_usuario') == 1
    for row in data:
        row['estadoLabel'] = 'Activo' if row['estado'] == 1 else 'Inactivo'
        if not es_admin:
            row['acciones'] = 'solo_lectura'
        else:
            row['acciones'] = 'editar,eliminar' if row['estado'] == 1 else 'reingresar'
    return jsonify(data)


@autor_bp.route('/registrar', methods=['POST'])
@permission_required('Autor')
def registrar():
    redir = login_required()
    if redir:
        return redir
    autor = str_clean(request.form.get('autor', ''))
    id_ = str_clean(request.form.get('id', ''))
    if not autor:
        return jsonify({'msg': 'El campo autor es requerido', 'icono': 'warning'})
    if not SecurityUtils.validate_person_name(autor, max_len=120):
        return jsonify({'msg': 'El nombre del autor debe tener solo letras y espacios (3-120 caracteres)', 'icono': 'warning'})
    img = request.files.get('imagen')
    url_imagen = str_clean(request.form.get('url_imagen', ''))
    if url_imagen and not SecurityUtils.validate_url(url_imagen):
        return jsonify({'msg': 'La URL de la imagen no es valida', 'icono': 'warning'})
    img_nombre = 'default.png'
    upload_folder = os.path.join(current_app.root_path, 'static', 'img', 'autores')
    os.makedirs(upload_folder, exist_ok=True)
    
    if img and img.filename:
        ext = img.filename.rsplit('.', 1)[-1].lower()
        is_image = img.content_type and img.content_type.startswith('image/')
        allowed_exts = ('png', 'jpg', 'jpeg', 'webp', 'gif', 'avif', 'svg', 'bmp', 'tiff', 'jfif', 'ico', 'heic', 'heif')
        if not is_image and ext not in allowed_exts:
            return jsonify({'msg': 'Archivo no permitido', 'icono': 'warning'})
        img_nombre = datetime.now().strftime('%Y%m%d%H%M%S') + f'.{ext}'
    elif url_imagen:
        import urllib.request
        from urllib.parse import urlparse
        try:
            parsed = urlparse(url_imagen)
            ext = 'jpg'
            if '.' in parsed.path:
                ext_guess = parsed.path.rsplit('.', 1)[-1].lower()
                is_image_guess = ext_guess in ('png', 'jpg', 'jpeg', 'webp', 'gif', 'avif', 'svg', 'bmp', 'tiff', 'jfif', 'ico', 'heic', 'heif')
                if is_image_guess:
                    ext = ext_guess
            img_nombre = datetime.now().strftime('%Y%m%d%H%M%S') + f'_url.{ext}'
            path_temp = os.path.join(upload_folder, img_nombre)
            req = urllib.request.Request(url_imagen, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=10) as response, open(path_temp, 'wb') as out_file:
                out_file.write(response.read())
        except Exception as e:
            return jsonify({'msg': 'No se pudo descargar la imagen de la URL', 'icono': 'warning'})
    elif request.form.get('foto_actual'):
        img_nombre = request.form.get('foto_actual')
    m = AutorModel()
    if id_ == '':
        data = m.insertar_autor(autor, img_nombre)
        if data == 'ok':
            if img and img.filename:
                os.makedirs(upload_folder, exist_ok=True)
                img.save(os.path.join(upload_folder, img_nombre))
            return jsonify({'msg': 'Autor registrado', 'icono': 'success'})
        elif data == 'existe':
            return jsonify({'msg': 'El autor ya existe', 'icono': 'warning'})
        return jsonify({'msg': 'Error al registrar', 'icono': 'error'})
    else:
        data = m.actualizar_autor(autor, img_nombre, id_)
        if data == 'modificado':
            if img and img.filename:
                os.makedirs(upload_folder, exist_ok=True)
                img.save(os.path.join(upload_folder, img_nombre))
            return jsonify({'msg': 'Autor modificado', 'icono': 'success'})
        return jsonify({'msg': 'Error al modificar', 'icono': 'error'})


@autor_bp.route('/editar/<int:id_>')
@permission_required('Autor')
def editar(id_):
    redir = login_required()
    if redir:
        return redir
    return jsonify(AutorModel().edit_autor(id_))


@autor_bp.route('/eliminar/<int:id_>')
@permission_required('Autor')
def eliminar(id_):
    redir = login_required()
    if redir:
        return redir
    data = AutorModel().estado_autor(0, id_)
    if data == 1:
        return jsonify({'msg': 'Autor dado de baja', 'icono': 'success'})
    return jsonify({'msg': 'Error al eliminar', 'icono': 'error'})


@autor_bp.route('/reingresar/<int:id_>')
@permission_required('Autor')
def reingresar(id_):
    redir = login_required()
    if redir:
        return redir
    data = AutorModel().estado_autor(1, id_)
    if data == 1:
        return jsonify({'msg': 'Autor restaurado', 'icono': 'success'})
    return jsonify({'msg': 'Error al restaurar', 'icono': 'error'})


@autor_bp.route('/buscarAutor')
@permission_required('Autor')
def buscar_autor():
    redir = login_required()
    if redir:
        return redir
    valor = request.args.get('au', '')
    return jsonify(AutorModel().buscar_autor(valor))


@autor_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir:
        return redir
    from fpdf import FPDF
    from flask import Response
    m = AutorModel()
    autores = m.get_autor()
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Autores')
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.cell(195, 8, 'Listado de Autores', 0, 1, 'C')
    pdf_doc.ln(3)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(6, 11, 20)
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(14, 6, 'N', 1, 0, 'C', True)
    pdf_doc.cell(150, 6, 'Autor', 1, 0, 'C', True)
    pdf_doc.cell(31, 6, 'Estado', 1, 1, 'C', True)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(autores, 1):
        estado = 'Activo' if row.get('estado') == 1 else 'Inactivo'
        pdf_doc.cell(14, 6, str(i), 1, 0, 'C')
        pdf_doc.cell(150, 6, str(row.get('autor', '')), 1, 0, 'L')
        pdf_doc.cell(31, 6, estado, 1, 1, 'C')
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=autores.pdf'})


@autor_bp.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir:
        return redir
    from io import BytesIO
    from datetime import date
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response
    m = AutorModel()
    autores = m.get_autor()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Autores"
    header_fill = PatternFill("solid", fgColor="060b14")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    for col, h in enumerate(['N°', 'Autor', 'Estado'], 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for i, row in enumerate(autores, 1):
        ws.append([i, row.get('autor', ''), 'Activo' if row.get('estado') == 1 else 'Inactivo'])
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=autores_{today}.xlsx'}
    )


@autor_bp.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir:
        return redir
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import Response
    m = AutorModel()
    autores = m.get_autor()
    doc = Document()
    title = doc.add_heading('Listado de Autores', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Generado el {date.today().isoformat()}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['N°', 'Autor', 'Estado']):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for i, row in enumerate(autores, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(row.get('autor', ''))
        cells[2].text = 'Activo' if row.get('estado') == 1 else 'Inactivo'
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=autores_{today}.docx'}
    )
