import sys
import os

# Bootstrapping para permitir ejecucion directa si es necesario
if __name__ == "__main__" or __package__ is None:
    current = os.path.dirname(os.path.realpath(__file__))
    parent = os.path.dirname(current)
    if parent not in sys.path:
        sys.path.append(parent)

from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template
from models.usuarios_model import UsuariosModel
from helpers import str_clean
from security import SecurityUtils, login_required, permission_required
from logger import get_logger
from exceptions import ValidationError, AuthenticationError

usuarios_bp = Blueprint('usuarios', __name__)
logger = get_logger('app')


@usuarios_bp.route('/')
@login_required
def index():
    try:
        id_user = session['id_usuario']
        m = UsuariosModel()
        perm = m.verificar_permisos(id_user, "Usuarios")
        if not perm and id_user != 1:
            return render_template('permisos.html')
        return render_template('usuarios/index.html')
    except Exception as e:
        logger.error(f"Error in usuarios index: {e}")
        return render_template('permisos.html')


@usuarios_bp.route('/listar')
@login_required
@permission_required('Usuarios')
def listar():
    try:
        id_caller = session.get('id_usuario')
        es_admin = (id_caller == 1)
        m = UsuariosModel()
        data = m.get_usuarios()
        for row in data:
            if not es_admin:
                # Usuario regular: solo ve info, sin botones de acción
                row['estadoLabel'] = 'Activo' if row['estado'] == 1 else 'Inactivo'
                row['acciones'] = 'solo_lectura'
            elif row['estado'] == 1:
                if row['id'] != 1:
                    row['estadoLabel'] = 'Activo'
                    row['acciones'] = 'editar,eliminar,roles'
                else:
                    row['estadoLabel'] = 'Activo'
                    row['acciones'] = 'superadmin'
            else:
                row['estadoLabel'] = 'Inactivo'
                row['acciones'] = 'reingresar'
        return jsonify(data)
    except Exception as e:
        logger.error(f"Error in listar: {e}")
        return jsonify({'error': 'Error al listar usuarios'}), 500


@usuarios_bp.route('/validar', methods=['POST'])
def validar():
    try:
        usuario = SecurityUtils.sanitize_input(request.form.get('usuario', ''))
        clave = request.form.get('clave', '')

        if not usuario or not clave:
            raise ValidationError('Todos los campos son requeridos')

        m = UsuariosModel()
        data = m.get_usuario(usuario, clave)
        if data:
            permisos = [row['nombre'] for row in m.get_permisos_usuario(data['id'])]
            session['id_usuario'] = data['id']
            session['usuario'] = data['usuario']
            session['nombre'] = data['nombre']
            session['activo'] = True
            session['permisos'] = permisos
            session['rol'] = data.get('rol', 'usuario')
            url_redir = '/Configuracion/admin'
            return jsonify({'msg': 'Procesando', 'icono': 'success', 'url': url_redir})

        logger.warning(f"Failed login attempt for user {usuario}")
        raise AuthenticationError('Usuario o contraseña incorrecta')

    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except AuthenticationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in validar: {e}")
        return jsonify({'msg': 'Error en autenticación', 'icono': 'error'}), 500


@usuarios_bp.route('/registro_publico', methods=['POST'])
def registro_publico():
    try:
        usuario = SecurityUtils.sanitize_input(request.form.get('usuario', ''))
        nombre = SecurityUtils.sanitize_input(request.form.get('nombre', ''))
        clave = request.form.get('clave', '')
        correo = SecurityUtils.sanitize_input(request.form.get('correo', ''))
        
        if not usuario or not nombre or not clave or not correo:
            raise ValidationError('Todos los campos son requeridos')
        if not SecurityUtils.validate_email(correo):
            raise ValidationError('Ingresa un correo electrónico válido')
        if not SecurityUtils.validate_username(usuario):
            raise ValidationError('El usuario solo admite letras, numeros, guiones y guion bajo (3-20 caracteres)')
        if not SecurityUtils.validate_person_name(nombre, max_len=100):
            raise ValidationError('El nombre debe tener solo letras y espacios (3-100 caracteres)')

        is_strong, msg = SecurityUtils.validate_password_strength(clave)
        if not is_strong:
            raise ValidationError(msg)

        m = UsuariosModel()
        data = m.registrar_usuario(usuario, nombre, clave, correo=correo)
        
        if data == 'existe':
            raise ValidationError('El usuario ya existe')
            
        if data == 'ok':
            # Assign limited permissions to this new user (1=Libros, 2=Autor, 3=Editorial, 7=Materias)
            user = m.get_usuario(usuario, clave)
            if user:
                for permiso_id in [1, 2, 3, 7]:
                    m.actualizar_permisos(user['id'], permiso_id)
                
                # Auto log in the user
                permisos = [row['nombre'] for row in m.get_permisos_usuario(user['id'])]
                session['id_usuario'] = user['id']
                session['usuario'] = user['usuario']
                session['nombre'] = user['nombre']
                session['activo'] = True
                session['permisos'] = permisos
                session['rol'] = user.get('rol', 'usuario')
                return jsonify({'msg': 'Registro exitoso, ingresando...', 'icono': 'success', 'url': '/Configuracion/admin'})
        
        return jsonify({'msg': 'Error al registrar', 'icono': 'error'})
    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in registro_publico: {e}")
        return jsonify({'msg': 'Error en el servidor', 'icono': 'error'}), 500


@usuarios_bp.route('/registrar', methods=['POST'])
@login_required
@permission_required('Usuarios')
def registrar():
    try:
        usuario = SecurityUtils.sanitize_input(request.form.get('usuario', ''))
        nombre = SecurityUtils.sanitize_input(request.form.get('nombre', ''))
        correo = SecurityUtils.sanitize_input(request.form.get('correo', ''))
        clave = request.form.get('clave', '')
        confirmar = request.form.get('confirmar', '')
        id_ = request.form.get('id', '')

        if not usuario or not nombre:
            raise ValidationError('Usuario y nombre son requeridos')
        if correo and not SecurityUtils.validate_email(correo):
            raise ValidationError('Ingresa un correo electrónico válido')
        if not SecurityUtils.validate_username(usuario):
            raise ValidationError('El usuario solo admite letras, numeros, guiones y guion bajo (3-20 caracteres)')
        if not SecurityUtils.validate_person_name(nombre, max_len=100):
            raise ValidationError('El nombre debe tener solo letras y espacios (3-100 caracteres)')

        m = UsuariosModel()

        if id_ == '':
            if not clave or not confirmar:
                raise ValidationError('La contraseña es requerida')
            if clave != confirmar:
                raise ValidationError('Las contraseñas no coinciden')

            is_strong, msg = SecurityUtils.validate_password_strength(clave)
            if not is_strong:
                raise ValidationError(msg)

            data = m.registrar_usuario(usuario, nombre, clave, correo=correo)
            msgs = {
                'ok': ('Usuario registrado', 'success'),
                'existe': ('El usuario ya existe', 'warning'),
                'error': ('Error al registrar', 'error')
            }
            msg, icono = msgs.get(data, ('Error', 'error'))
            logger.info(f"User {usuario} registered by {session.get('usuario')}")
            return jsonify({'msg': msg, 'icono': icono})
        else:
            data = m.modificar_usuario(usuario, nombre, id_, correo=correo)
            if data == 'existe':
                return jsonify({'msg': 'El usuario ya existe', 'icono': 'warning'})
            if data == 'modificado':
                logger.info(f"User {id_} modified by {session.get('usuario')}")
                return jsonify({'msg': 'Usuario modificado', 'icono': 'success'})
            return jsonify({'msg': 'Error al modificar', 'icono': 'error'})

    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in registrar: {e}")
        return jsonify({'msg': 'Error al registrar', 'icono': 'error'}), 500


@usuarios_bp.route('/editar/<int:id_>')
@login_required
@permission_required('Usuarios')
def editar(id_):
    try:
        m = UsuariosModel()
        result = m.editar_user(id_)
        if not result:
            return jsonify({'error': 'Usuario no encontrado'}), 404
        return jsonify(result)
    except Exception as e:
        logger.error(f"Error in editar: {e}")
        return jsonify({'error': 'Error al editar'}), 500


@usuarios_bp.route('/eliminar/<int:id_>')
@login_required
@permission_required('Usuarios')
def eliminar(id_):
    try:
        if id_ == 1:
            return jsonify({'msg': 'No puedes eliminar el admin', 'icono': 'warning'})

        m = UsuariosModel()
        data = m.accion_user(0, id_)
        if data == 1:
            logger.info(f"User {id_} deactivated by {session.get('usuario')}")
            return jsonify({'msg': 'Usuario dado de baja', 'icono': 'success'})
        return jsonify({'msg': 'Error al eliminar', 'icono': 'error'})
    except Exception as e:
        logger.error(f"Error in eliminar: {e}")
        return jsonify({'msg': 'Error al eliminar', 'icono': 'error'}), 500


@usuarios_bp.route('/reingresar/<int:id_>')
@login_required
@permission_required('Usuarios')
def reingresar(id_):
    try:
        m = UsuariosModel()
        data = m.accion_user(1, id_)
        if data == 1:
            logger.info(f"User {id_} reactivated by {session.get('usuario')}")
            return jsonify({'msg': 'Usuario restaurado', 'icono': 'success'})
        return jsonify({'msg': 'Error al restaurar', 'icono': 'error'})
    except Exception as e:
        logger.error(f"Error in reingresar: {e}")
        return jsonify({'msg': 'Error al restaurar', 'icono': 'error'}), 500


@usuarios_bp.route('/permisos/<int:id_>')
@login_required
@permission_required('Usuarios')
def permisos(id_):
    try:
        m = UsuariosModel()
        todos = m.get_permisos()
        asignados = m.get_detalle_permisos(id_)
        asignados_ids = {a['id_permiso'] for a in asignados}
        return jsonify({'permisos': todos, 'asignados': list(asignados_ids), 'id_usuario': id_})
    except Exception as e:
        logger.error(f"Error in permisos: {e}")
        return jsonify({'error': 'Error al obtener permisos'}), 500


@usuarios_bp.route('/registrarPermisos', methods=['POST'])
@login_required
@permission_required('Usuarios')
def registrar_permisos():
    try:
        id_user = request.form.get('id_usuario', '')
        permisos_list = request.form.getlist('permisos[]')

        if not id_user:
            raise ValidationError('ID de usuario requerido')

        m = UsuariosModel()
        m.delete_permisos(id_user)
        for permiso in permisos_list:
            m.actualizar_permisos(id_user, permiso)

        logger.info(f"Permissions updated for user {id_user} by {session.get('usuario')}")
        return jsonify({'msg': 'Permisos actualizados', 'icono': 'success'})
    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in registrar_permisos: {e}")
        return jsonify({'msg': 'Error al registrar permisos', 'icono': 'error'}), 500


@usuarios_bp.route('/cambiarPas', methods=['POST'])
@login_required
def cambiar_pas():
    try:
        id_ = session['id_usuario']
        clave_actual = request.form.get('clave_actual', '')
        clave_nueva = request.form.get('clave_nueva', '')
        confirmar = request.form.get('confirmar', '')

        if not clave_actual or not clave_nueva:
            raise ValidationError('Campos requeridos')
        if not confirmar:
            raise ValidationError('Debes confirmar la nueva contraseña')

        if clave_nueva != confirmar:
            raise ValidationError('Las contraseñas no coinciden')

        is_strong, msg = SecurityUtils.validate_password_strength(clave_nueva)
        if not is_strong:
            raise ValidationError(msg)

        m = UsuariosModel()
        user = m.editar_user(id_)

        if not user:
            raise ValidationError('Usuario no encontrado')

        if not SecurityUtils.verify_password(clave_actual, user['clave']):
            logger.warning(f"Invalid password change attempt for user {id_}")
            raise ValidationError('Contraseña actual incorrecta')

        data = m.actualizar_pass(clave_nueva, id_)
        if data == 'modificado':
            logger.info(f"Password changed for user {id_}")
            return jsonify({'msg': 'Contraseña modificada', 'icono': 'success'})
        return jsonify({'msg': 'Error al modificar', 'icono': 'warning'})

    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in cambiar_pas: {e}")
        return jsonify({'msg': 'Error al cambiar contraseña', 'icono': 'error'}), 500


@usuarios_bp.route('/salir')
def salir():
    try:
        usuario = session.get('usuario', 'desconocido')
        session.clear()
        logger.info(f"User {usuario} logged out")
        return redirect(url_for('landing'))
    except Exception as e:
        logger.error(f"Error in salir: {e}")
        return redirect(url_for('landing'))


@usuarios_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    if not session.get('activo'):
        return redirect(url_for('login'))
    from fpdf import FPDF
    from flask import Response
    m = UsuariosModel()
    registros = m.get_usuario()
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Listado de Usuarios')
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.cell(195, 8, 'Listado de Usuarios', 0, 1, 'C')
    pdf_doc.ln(3)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(26, 58, 92)
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(10, 6, 'N°', 1, 0, 'C', True)
    pdf_doc.cell(30, 6, 'Usuario', 1, 0, 'C', True)
    pdf_doc.cell(50, 6, 'Nombre', 1, 0, 'C', True)
    pdf_doc.cell(60, 6, 'Correo', 1, 0, 'C', True)
    pdf_doc.cell(30, 6, 'Estado', 1, 1, 'C', True)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(registros, 1):
        pdf_doc.cell(10, 6, str(i), 1, 0, 'C')
        pdf_doc.cell(30, 6, str(row.get('usuario', '')), 1, 0, 'L')
        pdf_doc.cell(50, 6, str(row.get('nombre', '')), 1, 0, 'L')
        pdf_doc.cell(60, 6, str(row.get('correo', '')), 1, 0, 'L')
        pdf_doc.cell(30, 6, 'Activo' if row.get('estado') == 1 else 'Inactivo', 1, 1, 'C')
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=reporte.pdf'})


@usuarios_bp.route('/excel')
@permission_required('Reportes')
def excel():
    if not session.get('activo'):
        return redirect(url_for('login'))
    from io import BytesIO
    from datetime import date
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response
    m = UsuariosModel()
    registros = m.get_usuario()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Listado de Usuarios"
    header_fill = PatternFill("solid", fgColor="1a3a5c")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers_arr = ['N°', 'Usuario', 'Nombre', 'Correo', 'Estado']
    for col, h in enumerate(headers_arr, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for i, row in enumerate(registros, 1):
        ws.append([i, row.get('usuario', ''), row.get('nombre', ''), row.get('correo', ''), 'Activo' if row.get('estado') == 1 else 'Inactivo'])
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.xlsx'}
    )


@usuarios_bp.route('/word')
@permission_required('Reportes')
def word():
    if not session.get('activo'):
        return redirect(url_for('login'))
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import Response
    m = UsuariosModel()
    registros = m.get_usuario()
    doc = Document()
    title = doc.add_heading('Listado de Usuarios', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Generado el {date.today().isoformat()}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    headers_arr = ['N°', 'Usuario', 'Nombre', 'Correo', 'Estado']
    table = doc.add_table(rows=1, cols=len(headers_arr))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers_arr):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs: run.bold = True
    for i, row in enumerate(registros, 1):
        cells = table.add_row().cells
        vals = [str(i), str(row.get('usuario', '')), str(row.get('nombre', '')), str(row.get('correo', '')), 'Activo' if row.get('estado') == 1 else 'Inactivo']
        for j, v in enumerate(vals):
            cells[j].text = v
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.docx'}
    )


@usuarios_bp.route('/mis_prestamos')
def mis_prestamos():
    if not session.get('activo'):
        return jsonify({'error': 'No autorizado'}), 401
    
    nombre_usuario = session.get('nombre')
    if not nombre_usuario:
        return jsonify([])
        
    from models.estudiantes_model import EstudiantesModel
    from models.prestamos_model import PrestamosModel
    
    em = EstudiantesModel()
    reg = em.select("SELECT id FROM estudiante WHERE nombre = %s LIMIT 1", (nombre_usuario,))
    if not reg:
        return jsonify([])
        
    pm = PrestamosModel()
    prestamos = pm.get_prestamos(estudiante=reg['id'])
    
    resultado = []
    for p in prestamos:
        resultado.append({
            'titulo': p.get('titulo', ''),
            'fecha_prestamo': str(p.get('fecha_prestamo', '')),
            'fecha_devolucion': str(p.get('fecha_devolucion', '')),
            'cantidad': p.get('cantidad', 1),
            'estado': p.get('estado', 1)
        })
        
    return jsonify(resultado)


@usuarios_bp.route('/solicitar_recuperacion', methods=['POST'])
def solicitar_recuperacion():
    try:
        identificador = SecurityUtils.sanitize_input(request.form.get('identificador', ''))
        if not identificador:
            raise ValidationError('Por favor ingresa tu usuario o correo electrónico')
            
        m = UsuariosModel()
        # Buscar usuario por nombre de usuario o por correo
        user = m.select("SELECT * FROM usuarios WHERE (usuario = %s OR correo = %s) AND estado = 1 LIMIT 1", (identificador, identificador))
        if not user:
            raise ValidationError('El usuario o correo electrónico no está registrado o está inactivo')
            
        # Generar un código aleatorio de 6 dígitos
        import random
        code = str(random.randint(100000, 999999))
        
        # Guardar en la base de datos
        m.save("UPDATE usuarios SET recovery_code = %s WHERE id = %s", (code, user['id']))
        
        logger.info(f"CÓDIGO DE RECUPERACIÓN GENERADO PARA {user['usuario']} ({user['correo']}): {code}")
        
        # Devolvemos el código en la respuesta para facilitar pruebas/desarrollo locales, además de simular envío por correo
        return jsonify({
            'msg': 'Código de recuperación generado con éxito. Revisa la consola o cópialo para restablecer tu clave.',
            'icono': 'success',
            'codigo_dev': code,
            'usuario': user['usuario']
        })
    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in solicitar_recuperacion: {e}")
        return jsonify({'msg': 'Error en el servidor', 'icono': 'error'}), 500


@usuarios_bp.route('/restablecer_clave', methods=['POST'])
def restablecer_clave():
    try:
        usuario = SecurityUtils.sanitize_input(request.form.get('usuario', ''))
        codigo = SecurityUtils.sanitize_input(request.form.get('codigo', ''))
        nueva_clave = request.form.get('nueva_clave', '')
        confirmar = request.form.get('confirmar', '')
        
        if not usuario or not codigo or not nueva_clave or not confirmar:
            raise ValidationError('Todos los campos son obligatorios')
            
        if nueva_clave != confirmar:
            raise ValidationError('Las contraseñas no coinciden')
            
        is_strong, msg = SecurityUtils.validate_password_strength(nueva_clave)
        if not is_strong:
            raise ValidationError(msg)
            
        m = UsuariosModel()
        user = m.select("SELECT * FROM usuarios WHERE usuario = %s AND estado = 1 LIMIT 1", (usuario,))
        if not user:
            raise ValidationError('Usuario no encontrado')
            
        if not user.get('recovery_code') or user['recovery_code'] != codigo:
            raise ValidationError('El código de recuperación es incorrecto o ha expirado')
            
        # Restablecer contraseña y borrar código
        hash_clave = SecurityUtils.hash_password(nueva_clave)
        m.save("UPDATE usuarios SET clave = %s, recovery_code = NULL WHERE id = %s", (hash_clave, user['id']))
        
        logger.info(f"Contraseña restablecida con éxito para el usuario {usuario}")
        return jsonify({'msg': 'Contraseña restablecida con éxito. Inicia sesión con tu nueva contraseña.', 'icono': 'success'})
    except ValidationError as e:
        return jsonify({'msg': e.message, 'icono': 'warning'}), e.code
    except Exception as e:
        logger.error(f"Error in restablecer_clave: {e}")
        return jsonify({'msg': 'Error en el servidor', 'icono': 'error'}), 500
