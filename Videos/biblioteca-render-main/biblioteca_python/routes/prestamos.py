from datetime import datetime, date
from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template, Response
from models.prestamos_model import PrestamosModel
from helpers import str_clean
from security import permission_required, SecurityUtils

prestamos_bp = Blueprint('prestamos', __name__)


def login_required():
    if not session.get('activo'):
        return redirect(url_for('login'))
    return None


@prestamos_bp.route('/')
def index():
    redir = login_required()
    if redir:
        return redir
    id_user = session['id_usuario']
    m = PrestamosModel()
    perm = m.verificar_permisos(id_user, "Prestamos")
    if not perm and id_user != 1:
        return render_template('permisos.html')
    return render_template('prestamos/index.html')


@prestamos_bp.route('/listar')
@permission_required('Prestamos')
def listar():
    redir = login_required()
    if redir:
        return redir
    estado       = request.args.get('estado')
    fecha_desde  = request.args.get('fecha_desde')
    fecha_hasta  = request.args.get('fecha_hasta')
    estudiante   = request.args.get('estudiante')
    libro        = request.args.get('libro')
    m = PrestamosModel()
    data = m.get_prestamos(estado=estado, fecha_desde=fecha_desde, fecha_hasta=fecha_hasta, estudiante=estudiante, libro=libro)
    for row in data:
        if row['estado'] == 1:
            row['estadoLabel'] = 'Prestado'
        elif row['estado'] == 0:
            row['estadoLabel'] = 'Entregado'
        elif row['estado'] == 2:
            row['estadoLabel'] = 'Pendiente'
        elif row['estado'] == 3:
            row['estadoLabel'] = 'Rechazado'
        else:
            row['estadoLabel'] = 'Desconocido'
            
        # Convertir fechas a string si son datetime
        for campo in ('fecha_prestamo', 'fecha_devolucion'):
            if row.get(campo) and not isinstance(row[campo], str):
                row[campo] = str(row[campo])
    return jsonify(data)


@prestamos_bp.route('/registrar', methods=['POST'])
@permission_required('Prestamos')
def registrar():
    redir = login_required()
    if redir:
        return redir
    estudiante = str_clean(request.form.get('estudiante', ''))
    libro = str_clean(request.form.get('libro', ''))
    cantidad = str_clean(request.form.get('cantidad', ''))
    fecha_prestamo = str_clean(request.form.get('fecha_prestamo', ''))
    fecha_devolucion = str_clean(request.form.get('fecha_devolucion', ''))
    observacion = str_clean(request.form.get('observacion', ''))

    if session.get('id_usuario') != 1:
        from models.estudiantes_model import EstudiantesModel
        em = EstudiantesModel()
        nombre_usuario = session.get('nombre', 'Usuario')
        reg = em.select("SELECT id FROM estudiante WHERE nombre = %s LIMIT 1", (nombre_usuario,))
        if not reg:
            codigo = "USR" + str(session.get('id_usuario'))
            em.insertar_estudiante(codigo, '00000000', nombre_usuario, 'Usuario Web', '-', '-')
            reg = em.select("SELECT id FROM estudiante WHERE nombre = %s LIMIT 1", (nombre_usuario,))
        if reg:
            estudiante = str(reg['id'])

    estado_ini = 1 if session.get('id_usuario') == 1 else 2
    
    if not estudiante or not libro or not cantidad or not fecha_prestamo or not fecha_devolucion:
        return jsonify({'msg': 'Todos los campos son requeridos', 'icono': 'warning'})
    if not SecurityUtils.validate_positive_int(estudiante):
        return jsonify({'msg': 'Selecciona un estudiante valido', 'icono': 'warning'})
    if not SecurityUtils.validate_positive_int(libro):
        return jsonify({'msg': 'Selecciona un libro valido', 'icono': 'warning'})
    if not SecurityUtils.validate_positive_int(cantidad):
        return jsonify({'msg': 'La cantidad debe ser mayor a cero', 'icono': 'warning'})
    try:
        fecha_prestamo_dt = datetime.strptime(fecha_prestamo, '%Y-%m-%d').date()
        fecha_devolucion_dt = datetime.strptime(fecha_devolucion, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'msg': 'Las fechas ingresadas no son validas', 'icono': 'warning'})
    if fecha_devolucion_dt <= fecha_prestamo_dt:
        return jsonify({'msg': 'La fecha de devolucion debe ser posterior a la fecha de prestamo', 'icono': 'warning'})
    m = PrestamosModel()
    data = m.insertar_prestamo(estudiante, libro, cantidad,
                               fecha_prestamo, fecha_devolucion, observacion, estado_inicial=estado_ini)
    if data == 'existe':
        return jsonify({'msg': 'Este estudiante ya tiene este libro prestado o una solicitud activa', 'icono': 'warning'})
    if data == 'sin_stock':
        return jsonify({'msg': 'Cantidad insuficiente en stock', 'icono': 'warning'})
    if data == 'cantidad_invalida':
        return jsonify({'msg': 'Cantidad invalida', 'icono': 'warning'})
    if data == 'libro_no_encontrado':
        return jsonify({'msg': 'Libro no encontrado', 'icono': 'warning'})
    if data and data > 0:
        msg = 'Prestamo registrado' if estado_ini == 1 else 'Solicitud enviada, en espera de aprobación'
        return jsonify({'msg': msg, 'icono': 'success', 'id': data})
    return jsonify({'msg': 'Error al registrar', 'icono': 'error'})


@prestamos_bp.route('/entregar/<int:id_>')
@permission_required('Prestamos')
def entregar(id_):
    redir = login_required()
    if redir:
        return redir
    if session.get('rol') not in ('administrador', 'bibliotecario') and session.get('id_usuario') != 1:
        return jsonify({'msg': 'No tienes permisos para realizar esta acción', 'icono': 'warning'}), 403
    m = PrestamosModel()
    data = m.actualizar_prestamo(0, id_)
    if data == 'ok':
        return jsonify({'msg': 'Libro entregado', 'icono': 'success'})
    return jsonify({'msg': 'Error al registrar entrega', 'icono': 'error'})

@prestamos_bp.route('/aprobar/<int:id_>')
@permission_required('Prestamos')
def aprobar(id_):
    redir = login_required()
    if redir:
        return redir
    if session.get('rol') not in ('administrador', 'bibliotecario') and session.get('id_usuario') != 1:
        return jsonify({'msg': 'No tienes permisos para realizar esta acción', 'icono': 'warning'}), 403
    m = PrestamosModel()
    data = m.aprobar_prestamo(id_)
    if data == 'ok':
        return jsonify({'msg': 'Préstamo aprobado. Stock actualizado.', 'icono': 'success'})
    elif data == 'sin_stock':
        return jsonify({'msg': 'No hay stock suficiente para aprobar.', 'icono': 'warning'})
    return jsonify({'msg': 'Error al aprobar préstamo', 'icono': 'error'})

@prestamos_bp.route('/pendientes_notificacion')
def pendientes_notificacion():
    redir = login_required()
    if redir:
        return jsonify([])
    
    # Mejora: Asegurar comparación de enteros para evitar errores de tipo en sesión
    es_admin = int(session.get('id_usuario', 0)) == 1
    tiene_permiso = 'Prestamos' in session.get('permisos', [])
    
    if not es_admin and not tiene_permiso:
        return jsonify([])
        
    m = PrestamosModel()
    all_data = m.get_prestamos()
    pendientes = [p for p in all_data if p['estado'] == 2]
    return jsonify(pendientes)

@prestamos_bp.route('/rechazar/<int:id_>')
@permission_required('Prestamos')
def rechazar(id_):
    redir = login_required()
    if redir:
        return redir
    if session.get('rol') not in ('administrador', 'bibliotecario') and session.get('id_usuario') != 1:
        return jsonify({'msg': 'No tienes permisos para realizar esta acción', 'icono': 'warning'}), 403
    m = PrestamosModel()
    data = m.rechazar_prestamo(id_)
    if data == 'ok':
        return jsonify({'msg': 'Préstamo rechazado.', 'icono': 'success'})
    return jsonify({'msg': 'Error al rechazar préstamo', 'icono': 'error'})

@prestamos_bp.route('/mis_notificaciones')
def mis_notificaciones():
    redir = login_required()
    if redir:
        return jsonify([])
    
    # Obtener el ID del estudiante para este usuario regular
    if int(session.get('id_usuario', 0)) == 1:
        return jsonify([]) # Admin usa pendientes_notificacion
        
    from models.estudiantes_model import EstudiantesModel
    em = EstudiantesModel()
    nombre_usuario = session.get('nombre', 'Usuario')
    reg = em.select("SELECT id FROM estudiante WHERE nombre = %s LIMIT 1", (nombre_usuario,))
    
    if not reg:
        return jsonify([])
        
    estudiante_id = reg['id']
    m = PrestamosModel()
    all_data = m.get_prestamos()
    
    # Filtrar solo aprobados activos (1) y rechazados (3) de este estudiante. Se oculta el pendiente (2).
    mis_prestamos = []
    for p in all_data:
        if p['id_estudiante'] == estudiante_id and p['estado'] in [1, 3]:
            mis_prestamos.append(p)
            
    return jsonify(mis_prestamos)



@prestamos_bp.route('/ticked/<int:id_>')
@permission_required('Prestamos')
def ticked(id_):
    redir = login_required()
    if redir:
        return redir
    m = PrestamosModel()
    data = m.get_prestamo_libro(id_)
    if data:
        for campo in ('fecha_prestamo', 'fecha_devolucion'):
            if data.get(campo) and not isinstance(data[campo], str):
                data[campo] = str(data[campo])
    return jsonify(data)


@prestamos_bp.route('/ticket_pdf/<int:id_>')
@permission_required('Prestamos')
def ticket_pdf(id_):
    redir = login_required()
    if redir:
        return redir
    
    from models.configuracion_model import ConfiguracionModel
    from models.prestamos_model import PrestamosModel
    from fpdf import FPDF
    
    m = PrestamosModel()
    loan = m.get_prestamo_libro(id_)
    if not loan:
        return "Préstamo no encontrado", 404
        
    config_m = ConfiguracionModel()
    datos_empresa = config_m.select_configuracion() or {}
    
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(15, 15, 15)
    pdf_doc.set_title(f"Ticket Prestamo #{loan.get('pres_id')}")
    
    # Draw double border around page (margin 12mm)
    pdf_doc.set_draw_color(26, 58, 92) # Navy Blue
    pdf_doc.set_line_width(0.8)
    pdf_doc.rect(12, 12, 191.9, 255.4)
    
    pdf_doc.set_line_width(0.2)
    pdf_doc.rect(13, 13, 189.9, 253.4)
    
    def clean(val):
        if val is None:
            return ""
        return str(val).encode('latin-1', 'replace').decode('latin-1')
        
    # Header logo/Name
    pdf_doc.ln(10)
    pdf_doc.set_font('Arial', 'B', 16)
    pdf_doc.set_text_color(26, 58, 92)
    pdf_doc.cell(185, 8, clean(datos_empresa.get('nombre', 'Biblioteca')), 0, 1, 'C')
    
    # Subtitle
    pdf_doc.set_font('Arial', 'I', 10)
    pdf_doc.set_text_color(100, 100, 100)
    pdf_doc.cell(185, 5, "Comprobante Oficial de Prestamo", 0, 1, 'C')
    pdf_doc.ln(2)
    
    # Contact Info
    pdf_doc.set_font('Arial', '', 9)
    info_parts = []
    if datos_empresa.get('direccion'):
        info_parts.append(f"Direccion: {datos_empresa.get('direccion')}")
    if datos_empresa.get('telefono'):
        info_parts.append(f"Tel: {datos_empresa.get('telefono')}")
    if datos_empresa.get('correo'):
        info_parts.append(f"Email: {datos_empresa.get('correo')}")
        
    info_txt = "  |  ".join(info_parts)
    pdf_doc.cell(185, 5, clean(info_txt), 0, 1, 'C')
    pdf_doc.ln(6)
    
    # Separator Line
    pdf_doc.set_draw_color(180, 180, 180)
    pdf_doc.line(20, 50, 195, 50)
    pdf_doc.ln(8)
    
    # Title of document
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.cell(185, 6, "TICKET DE PRESTAMO DE LIBRO", 0, 1, 'C')
    
    pdf_doc.set_font('Arial', 'B', 11)
    pdf_doc.set_text_color(200, 50, 50)
    pdf_doc.cell(185, 6, f"Ticket No: {loan.get('pres_id'):06d}", 0, 1, 'C')
    pdf_doc.ln(8)
    
    # Detail table header
    pdf_doc.set_fill_color(240, 244, 248)
    pdf_doc.set_text_color(26, 58, 92)
    pdf_doc.set_font('Arial', 'B', 11)
    pdf_doc.cell(175, 7, "  DATOS DEL DETALLE", 0, 1, 'L', True)
    pdf_doc.ln(2)
    
    def draw_row(label, value):
        pdf_doc.set_font('Arial', 'B', 10)
        pdf_doc.set_text_color(50, 50, 50)
        pdf_doc.cell(45, 6, "  " + clean(label) + ":", 0, 0, 'L')
        pdf_doc.set_font('Arial', '', 10)
        pdf_doc.set_text_color(0, 0, 0)
        pdf_doc.cell(130, 6, clean(value), 0, 1, 'L')
        
    draw_row("Estudiante", loan.get('est_nombre', ''))
    draw_row("Codigo / Matricula", loan.get('codigo', '-'))
    draw_row("Carrera", loan.get('carrera', '-'))
    draw_row("Libro Solicitado", loan.get('titulo', ''))
    draw_row("Cantidad Prestada", str(loan.get('cantidad', 1)))
    draw_row("Observaciones", loan.get('observacion', '-') or '-')
    
    estados_label = {0: 'Entregado (Devuelto)', 1: 'Prestado (Activo)', 2: 'Pendiente de Aprobacion', 3: 'Rechazado'}
    estado_val = estados_label.get(loan.get('estado'), 'Desconocido')
    draw_row("Estado del Prestamo", estado_val)
    
    pdf_doc.ln(8)
    
    # Dates Box
    pdf_doc.set_fill_color(255, 243, 205) # Light yellow
    pdf_doc.set_draw_color(255, 238, 186)
    pdf_doc.set_line_width(0.5)
    pdf_doc.rect(20, 142, 175, 28, 'DF')
    
    pdf_doc.set_y(145)
    pdf_doc.set_text_color(133, 100, 4)
    pdf_doc.set_font('Arial', 'B', 10.5)
    pdf_doc.cell(175, 5, "FECHAS DE CONTROL DE PRESTAMO", 0, 1, 'C')
    pdf_doc.ln(1)
    
    # Columns for dates
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_text_color(50, 50, 50)
    pdf_doc.cell(87, 5, "Fecha de Recibo (Prestamo):", 0, 0, 'C')
    pdf_doc.cell(88, 5, "Fecha Limite de Devolucion:", 0, 1, 'C')
    
    # Format dates
    fp = loan.get('fecha_prestamo')
    fd = loan.get('fecha_devolucion')
    if isinstance(fp, (date, datetime)):
        fp = fp.strftime('%d/%m/%Y')
    elif isinstance(fp, str):
        try: fp = datetime.strptime(fp.split()[0], '%Y-%m-%d').strftime('%d/%m/%Y')
        except: pass
        
    if isinstance(fd, (date, datetime)):
        fd = fd.strftime('%d/%m/%Y')
    elif isinstance(fd, str):
        try: fd = datetime.strptime(fd.split()[0], '%Y-%m-%d').strftime('%d/%m/%Y')
        except: pass
        
    pdf_doc.set_font('Arial', 'B', 12)
    pdf_doc.set_text_color(0, 102, 204) # Blue
    pdf_doc.cell(87, 6, clean(fp), 0, 0, 'C')
    pdf_doc.set_text_color(204, 0, 0) # Red
    pdf_doc.cell(88, 6, clean(fd), 0, 1, 'C')
    
    pdf_doc.ln(25)
    
    # Signatures
    pdf_doc.set_y(200)
    pdf_doc.set_draw_color(150, 150, 150)
    pdf_doc.set_line_width(0.3)
    pdf_doc.line(30, 220, 90, 220)
    pdf_doc.line(125, 220, 185, 220)
    
    pdf_doc.set_y(222)
    pdf_doc.set_font('Arial', '', 9)
    pdf_doc.set_text_color(100, 100, 100)
    pdf_doc.cell(87, 4, "Firma del Estudiante", 0, 0, 'C')
    pdf_doc.cell(88, 4, "Firma de Biblioteca / Sello", 0, 1, 'C')
    
    # Footer Note
    pdf_doc.set_y(245)
    pdf_doc.set_font('Arial', 'I', 8.5)
    pdf_doc.set_text_color(120, 120, 120)
    msg_footer = "Nota: Conserve este comprobante. El retraso en la devolucion del material bibliografico\ngenerara sanciones de acuerdo al reglamento interno de la biblioteca."
    pdf_doc.multi_cell(175, 4, clean(msg_footer), 0, 'C')
    
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': f'inline; filename=ticket_prestamo_{id_}.pdf'})



@prestamos_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir:
        return redir
    from models.configuracion_model import ConfiguracionModel
    from fpdf import FPDF
    from models.prestamos_model import PrestamosModel
    estado      = request.args.get('estado')
    fecha_desde = request.args.get('fecha_desde')
    fecha_hasta = request.args.get('fecha_hasta')
    estudiante  = request.args.get('estudiante')
    libro       = request.args.get('libro')
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    today = date.today().isoformat()
    prestamo = PrestamosModel().get_prestamos(estado=estado, fecha_desde=fecha_desde, fecha_hasta=fecha_hasta, estudiante=estudiante, libro=libro)
    # Etiqueta de filtros
    filtros_txt = []
    estados_map = {'0': 'Entregados', '1': 'Prestados', '2': 'Pendientes', '3': 'Rechazados'}
    if estado is not None and estado != '': filtros_txt.append(estados_map.get(str(estado), f'Estado:{estado}'))
    if fecha_desde: filtros_txt.append(f'Desde:{fecha_desde}')
    if fecha_hasta: filtros_txt.append(f'Hasta:{fecha_hasta}')
    subtitulo = 'Filtros: ' + ', '.join(filtros_txt) if filtros_txt else 'Todos los registros'
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Prestamos')
    pdf_doc.set_font('Arial', 'B', 12)
    pdf_doc.cell(195, 5, datos.get('nombre', ''), 0, 1, 'C')
    pdf_doc.set_font('Arial', 'I', 9)
    pdf_doc.set_text_color(100, 100, 100)
    pdf_doc.cell(195, 5, subtitulo, 0, 1, 'C')
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.cell(20, 5, 'Telefono: ', 0, 0, 'L')
    pdf_doc.set_font('Arial', '', 10)
    pdf_doc.cell(20, 5, datos.get('telefono', ''), 0, 1, 'L')
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.cell(20, 5, 'Correo: ', 0, 0, 'L')
    pdf_doc.set_font('Arial', '', 10)
    pdf_doc.cell(20, 5, datos.get('correo', ''), 0, 1, 'L')
    pdf_doc.ln()
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(10, 22, 40) # Elegant deep navy blue
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(196, 6, 'Detalle de Préstamos', 1, 1, 'C', True)
    
    # Table headers with navy background
    pdf_doc.set_fill_color(20, 40, 75)
    pdf_doc.cell(14, 6, 'N', 1, 0, 'C', True)
    pdf_doc.cell(50, 6, 'Usuarios', 1, 0, 'L', True)
    pdf_doc.cell(72, 6, 'Libros', 1, 0, 'L', True)
    pdf_doc.cell(30, 6, 'Fecha Préstamo', 1, 0, 'C', True)
    pdf_doc.cell(15, 6, 'Cant.', 1, 0, 'C', True)
    pdf_doc.cell(15, 6, 'Estado', 1, 1, 'C', True)
    
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 9.5)
    estados_label = {0: 'Entregado', 1: 'Prestado', 2: 'Pendiente', 3: 'Rechazado'}
    for i, row in enumerate(prestamo, 1):
        bg_fill = (i % 2 == 0)
        if bg_fill:
            pdf_doc.set_fill_color(242, 245, 248) # Light blue-grey zebra striping
        else:
            pdf_doc.set_fill_color(255, 255, 255)
            
        pdf_doc.cell(14, 6, str(i), 1, 0, 'C', True)
        pdf_doc.cell(50, 6, str(row.get('est_nombre', '')), 1, 0, 'L', True)
        pdf_doc.cell(72, 6, str(row.get('titulo', '')), 1, 0, 'L', True)
        pdf_doc.cell(30, 6, str(row.get('fecha_prestamo', '')), 1, 0, 'C', True)
        pdf_doc.cell(15, 6, str(row.get('cantidad', '')), 1, 0, 'C', True)
        pdf_doc.cell(15, 6, estados_label.get(row.get('estado'), ''), 1, 1, 'C', True)
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=prestamos.pdf'})


@prestamos_bp.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir:
        return redir
    from models.configuracion_model import ConfiguracionModel
    from io import BytesIO
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from models.prestamos_model import PrestamosModel
    estado      = request.args.get('estado')
    fecha_desde = request.args.get('fecha_desde')
    fecha_hasta = request.args.get('fecha_hasta')
    estudiante  = request.args.get('estudiante')
    libro       = request.args.get('libro')
    today = date.today().isoformat()
    prestamo = PrestamosModel().get_prestamos(estado=estado, fecha_desde=fecha_desde, fecha_hasta=fecha_hasta, estudiante=estudiante, libro=libro) or []
    filtros_txt = []
    estados_map = {'0': 'Entregados', '1': 'Prestados', '2': 'Pendientes', '3': 'Rechazados'}
    if estado is not None and estado != '': filtros_txt.append(estados_map.get(str(estado), f'Estado:{estado}'))
    if fecha_desde: filtros_txt.append(f'Desde:{fecha_desde}')
    if fecha_hasta: filtros_txt.append(f'Hasta:{fecha_hasta}')
    subtitulo = 'Filtros: ' + ', '.join(filtros_txt) if filtros_txt else 'Todos los registros'
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Préstamos"
    ws.merge_cells('A1:G1')
    sub_cell = ws['A1']
    sub_cell.value = subtitulo
    sub_cell.font = Font(italic=True, color='FF555555', size=10)
    sub_cell.alignment = Alignment(horizontal='center')
    # Cabecera
    header_fill = PatternFill("solid", fgColor="060b14")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers = ['N°', 'Estudiante', 'Libro', 'Fecha Préstamo', 'Fecha Devolución', 'Cant.', 'Estado']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    estados_label = {0: 'Entregado', 1: 'Prestado', 2: 'Pendiente', 3: 'Rechazado'}
    # Datos
    for i, row in enumerate(prestamo, 1):
        ws.append([
            i,
            row.get('est_nombre', ''),
            row.get('titulo', ''),
            str(row.get('fecha_prestamo', '')),
            str(row.get('fecha_devolucion', '')),
            row.get('cantidad', ''),
            estados_label.get(row.get('estado'), '')
        ])
    # Ancho de columnas
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=prestamos_{today}.xlsx'}
    )


@prestamos_bp.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir:
        return redir
    from models.configuracion_model import ConfiguracionModel
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from models.prestamos_model import PrestamosModel
    estado      = request.args.get('estado')
    fecha_desde = request.args.get('fecha_desde')
    fecha_hasta = request.args.get('fecha_hasta')
    estudiante  = request.args.get('estudiante')
    libro       = request.args.get('libro')
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    today = date.today().isoformat()
    prestamo = PrestamosModel().get_prestamos(estado=estado, fecha_desde=fecha_desde, fecha_hasta=fecha_hasta, estudiante=estudiante, libro=libro) or []
    filtros_txt = []
    estados_map = {'0': 'Entregados', '1': 'Prestados', '2': 'Pendientes', '3': 'Rechazados'}
    if estado is not None and estado != '': filtros_txt.append(estados_map.get(str(estado), f'Estado:{estado}'))
    if fecha_desde: filtros_txt.append(f'Desde:{fecha_desde}')
    if fecha_hasta: filtros_txt.append(f'Hasta:{fecha_hasta}')
    subtitulo = 'Filtros: ' + ', '.join(filtros_txt) if filtros_txt else 'Todos los registros'
    doc = Document()
    # Título
    title = doc.add_heading(datos.get('nombre', 'Biblioteca'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Reporte de Préstamos — {today} | {subtitulo}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    estados_label = {0: 'Entregado', 1: 'Prestado', 2: 'Pendiente', 3: 'Rechazado'}
    # Tabla
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['N°', 'Estudiante', 'Libro', 'F. Préstamo', 'Cant.', 'Estado']):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for i, row in enumerate(prestamo, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row.get('est_nombre', '')
        cells[2].text = row.get('titulo', '')
        cells[3].text = str(row.get('fecha_prestamo', ''))
        cells[4].text = str(row.get('cantidad', ''))
        cells[5].text = estados_label.get(row.get('estado'), '')
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=prestamos_{today}.docx'}
    )
