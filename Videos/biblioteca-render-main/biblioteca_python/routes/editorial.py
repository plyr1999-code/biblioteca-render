from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template
from models.editorial_model import EditorialModel
from helpers import str_clean
from security import permission_required, SecurityUtils

editorial_bp = Blueprint('editorial', __name__)


def login_required():
    if not session.get('activo'):
        return redirect(url_for('index'))
    return None


@editorial_bp.route('/')
def index():
    redir = login_required()
    if redir:
        return redir
    id_user = session['id_usuario']
    m = EditorialModel()
    perm = m.verificar_permisos(id_user, "Editorial")
    if not perm and id_user != 1:
        return render_template('permisos.html')
    return render_template('editorial/index.html')


@editorial_bp.route('/listar')
@permission_required('Editorial')
def listar():
    redir = login_required()
    if redir:
        return redir
    estado = request.args.get('estado')
    data = EditorialModel().get_editorial()
    if estado is not None and estado != '':
        data = [r for r in data if str(r.get('estado', '')) == str(estado)]
    es_admin = session.get('id_usuario') == 1
    for row in data:
        row['estadoLabel'] = 'Activo' if row['estado'] == 1 else 'Inactivo'
        if not es_admin:
            row['acciones'] = 'solo_lectura'
        else:
            row['acciones'] = 'editar,eliminar' if row['estado'] == 1 else 'reingresar'
    return jsonify(data)


@editorial_bp.route('/registrar', methods=['POST'])
@permission_required('Editorial')
def registrar():
    redir = login_required()
    if redir:
        return redir
    editorial = str_clean(request.form.get('editorial', ''))
    id_ = str_clean(request.form.get('id', ''))
    if not editorial:
        return jsonify({'msg': 'El campo editorial es requerido', 'icono': 'warning'})
    if not SecurityUtils.validate_generic_text(editorial, min_len=2, max_len=100):
        return jsonify({'msg': 'Ingresa una editorial valida (2-100 caracteres)', 'icono': 'warning'})
    m = EditorialModel()
    if id_ == '':
        data = m.insertar_editorial(editorial)
        msgs = {'ok': ('Editorial registrada', 'success'),
                'existe': ('La editorial ya existe', 'warning'),
                'error': ('Error al registrar', 'error')}
        msg, icono = msgs.get(data, ('Error', 'error'))
        return jsonify({'msg': msg, 'icono': icono})
    else:
        data = m.actualizar_editorial(editorial, id_)
        if data == 'modificado':
            return jsonify({'msg': 'Editorial modificada', 'icono': 'success'})
        return jsonify({'msg': 'Error al modificar', 'icono': 'error'})


@editorial_bp.route('/editar/<int:id_>')
@permission_required('Editorial')
def editar(id_):
    redir = login_required()
    if redir:
        return redir
    return jsonify(EditorialModel().edit_editorial(id_))


@editorial_bp.route('/eliminar/<int:id_>')
@permission_required('Editorial')
def eliminar(id_):
    redir = login_required()
    if redir:
        return redir
    data = EditorialModel().estado_editorial(0, id_)
    if data == 1:
        return jsonify({'msg': 'Editorial dada de baja', 'icono': 'success'})
    return jsonify({'msg': 'Error al eliminar', 'icono': 'error'})


@editorial_bp.route('/reingresar/<int:id_>')
@permission_required('Editorial')
def reingresar(id_):
    redir = login_required()
    if redir:
        return redir
    data = EditorialModel().estado_editorial(1, id_)
    if data == 1:
        return jsonify({'msg': 'Editorial restaurada', 'icono': 'success'})
    return jsonify({'msg': 'Error al restaurar', 'icono': 'error'})


@editorial_bp.route('/buscarEditorial')
@permission_required('Editorial')
def buscar_editorial():
    redir = login_required()
    if redir:
        return redir
    valor = request.args.get('ed', '')
    return jsonify(EditorialModel().buscar_editorial(valor))


@editorial_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir: return redir
    from fpdf import FPDF
    from flask import Response
    estado = request.args.get('estado')
    m = EditorialModel()
    registros = m.get_editorial()
    if estado is not None and estado != '':
        registros = [r for r in registros if str(r.get('estado', '')) == str(estado)]
    filtro_lbl = ('Activos' if str(estado)=='1' else 'Inactivos') if (estado is not None and estado != '') else 'Todos'
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Listado de Editoriales')
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.cell(195, 8, 'Listado de Editoriales', 0, 1, 'C')
    pdf_doc.set_font('Arial', 'I', 9)
    pdf_doc.set_text_color(100, 100, 100)
    pdf_doc.cell(195, 5, f'Estado: {filtro_lbl}', 0, 1, 'C')
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.ln(3)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(6, 11, 20)
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(15, 6, 'N°', 1, 0, 'C', True)
    pdf_doc.cell(130, 6, 'Editorial', 1, 0, 'C', True)
    pdf_doc.cell(30, 6, 'Estado', 1, 1, 'C', True)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(registros, 1):
        pdf_doc.cell(15, 6, str(i), 1, 0, 'C')
        pdf_doc.cell(130, 6, str(row.get('editorial', '')), 1, 0, 'L')
        pdf_doc.cell(30, 6, 'Activo' if row.get('estado') == 1 else 'Inactivo', 1, 1, 'C')
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=reporte.pdf'})


@editorial_bp.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response
    estado = request.args.get('estado')
    m = EditorialModel()
    registros = m.get_editorial()
    if estado is not None and estado != '':
        registros = [r for r in registros if str(r.get('estado', '')) == str(estado)]
    filtro_lbl = ('Activos' if str(estado)=='1' else 'Inactivos') if (estado is not None and estado != '') else 'Todos'
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Listado de Editoriales"
    ws.merge_cells('A1:C1')
    sub_cell = ws['A1']
    sub_cell.value = f'Estado: {filtro_lbl}'
    sub_cell.font = Font(italic=True, color='FF555555', size=10)
    sub_cell.alignment = Alignment(horizontal='center')
    header_fill = PatternFill("solid", fgColor="060b14")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers_arr = ['N°', 'Editorial', 'Estado']
    for col, h in enumerate(headers_arr, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for i, row in enumerate(registros, 1):
        ws.append([i, row.get('editorial', ''), 'Activo' if row.get('estado') == 1 else 'Inactivo'])
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.xlsx'}
    )


@editorial_bp.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import Response
    estado = request.args.get('estado')
    m = EditorialModel()
    registros = m.get_editorial()
    if estado is not None and estado != '':
        registros = [r for r in registros if str(r.get('estado', '')) == str(estado)]
    filtro_lbl = ('Activos' if str(estado)=='1' else 'Inactivos') if (estado is not None and estado != '') else 'Todos'
    doc = Document()
    title = doc.add_heading('Listado de Editoriales', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Generado el {date.today().isoformat()} | Estado: {filtro_lbl}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    headers_arr = ['N°', 'Editorial', 'Estado']
    table = doc.add_table(rows=1, cols=len(headers_arr))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers_arr):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs: run.bold = True
    for i, row in enumerate(registros, 1):
        cells = table.add_row().cells
        vals = [str(i), str(row.get('editorial', '')), 'Activo' if row.get('estado') == 1 else 'Inactivo']
        for j, v in enumerate(vals):
            cells[j].text = v
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.docx'}
    )
