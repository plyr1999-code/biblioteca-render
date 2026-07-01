import os
import re
from datetime import datetime
import sys
import os

# Bootstrapping para permitir ejecucion directa si es necesario
if __name__ == "__main__" or __package__ is None:
    current = os.path.dirname(os.path.realpath(__file__))
    parent = os.path.dirname(current)
    if parent not in sys.path:
        sys.path.append(parent)

from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template, current_app
from models.libros_model import LibrosModel
from helpers import str_clean
from security import permission_required, SecurityUtils

libros_bp = Blueprint('libros', __name__)

def normalize_anio_edicion(value: str) -> str:
    value = value.strip()
    if not value:
        return ''
    if re.fullmatch(r'\d{4}', value):
        return f"{value}-01-01"
    if re.fullmatch(r'\d{4}-\d{2}-\d{2}', value):
        return value
    return ''


def login_required():
    if not session.get('activo'):
        return redirect(url_for('login'))
    return None


@libros_bp.route('/')
def index():
    redir = login_required()
    if redir:
        return redir
    id_user = session['id_usuario']
    m = LibrosModel()
    perm = m.verificar_permisos(id_user, "Libros")
    if not perm and id_user != 1:
        return render_template('permisos.html')
    return render_template('libros/index.html')


@libros_bp.route('/listar')
@permission_required('Libros')
def listar():
    redir = login_required()
    if redir:
        return redir
    es_admin = int(session.get('id_usuario', 0)) == 1
    
    id_autor = request.args.get('id_autor') or None
    id_editorial = request.args.get('id_editorial') or None
    id_materia = request.args.get('id_materia') or None
    estado = request.args.get('estado') or None
    stock = request.args.get('stock') or None
    q = request.args.get('q', '')
    
    m = LibrosModel()
    # Realizamos la búsqueda directamente en la base de datos para mayor eficiencia
    data = m.get_libros(id_autor=id_autor, id_editorial=id_editorial, id_materia=id_materia, estado=estado, stock=stock, q=q)
    
    for row in data:
        row['estadoLabel'] = 'Activo' if row['estado'] == 1 else 'Inactivo'
        row['acciones'] = 'editar,eliminar' if row['estado'] == 1 else 'reingresar'
        row['es_admin'] = es_admin
    return jsonify(data)


@libros_bp.route('/registrar', methods=['POST'])
@permission_required('Libros')
def registrar():
    redir = login_required()
    if redir:
        return redir
    titulo = str_clean(request.form.get('titulo', ''))
    autor = str_clean(request.form.get('autor', ''))
    editorial = str_clean(request.form.get('editorial', ''))
    materia = str_clean(request.form.get('materia', ''))
    cantidad = str_clean(request.form.get('cantidad', ''))
    num_pagina = str_clean(request.form.get('num_pagina', ''))
    anio_edicion = str_clean(request.form.get('anio_edicion', ''))
    descripcion = str_clean(request.form.get('descripcion', ''))
    id_ = str_clean(request.form.get('id', ''))
    current_year = datetime.now().year

    if not titulo or not autor or not editorial or not materia or not cantidad or not num_pagina or not anio_edicion:
        return jsonify({'msg': 'Todos los campos son requeridos', 'icono': 'warning'})
    if not SecurityUtils.validate_generic_text(titulo, min_len=2, max_len=150):
        return jsonify({'msg': 'Ingresa un titulo valido (2-150 caracteres)', 'icono': 'warning'})

    m = LibrosModel()

    # Procesar autor (puede ser ID o nuevo nombre)
    if autor.isdigit():
        id_autor = int(autor)
    else:
        if not SecurityUtils.validate_person_name(autor, max_len=120):
            return jsonify({'msg': 'El nombre del autor debe tener solo letras y espacios (3-120 caracteres)', 'icono': 'warning'})
        id_autor = m.get_or_create_autor(autor)
        if not id_autor:
            return jsonify({'msg': 'Error al registrar el nuevo autor', 'icono': 'error'})

    # Procesar editorial (puede ser ID o nuevo nombre)
    if editorial.isdigit():
        id_editorial = int(editorial)
    else:
        if not SecurityUtils.validate_generic_text(editorial, min_len=2, max_len=100):
            return jsonify({'msg': 'Ingresa una editorial valida (2-100 caracteres)', 'icono': 'warning'})
        id_editorial = m.get_or_create_editorial(editorial)
        if not id_editorial:
            return jsonify({'msg': 'Error al registrar la nueva editorial', 'icono': 'error'})

    # Procesar materia (puede ser ID o nuevo nombre)
    if materia.isdigit():
        id_materia = int(materia)
    else:
        if not SecurityUtils.validate_generic_text(materia, min_len=2, max_len=100):
            return jsonify({'msg': 'Ingresa una materia valida (2-100 caracteres)', 'icono': 'warning'})
        id_materia = m.get_or_create_materia(materia)
        if not id_materia:
            return jsonify({'msg': 'Error al registrar la nueva materia', 'icono': 'error'})

    if descripcion and not SecurityUtils.validate_generic_text(descripcion, min_len=3, max_len=500):
        return jsonify({'msg': 'La descripcion contiene caracteres no permitidos o supera los 500 caracteres', 'icono': 'warning'})
    anio_edicion = normalize_anio_edicion(anio_edicion)
    if not anio_edicion:
        return jsonify({'msg': 'Anio de edicion invalido. Usa AAAA o AAAA-MM-DD', 'icono': 'warning'})
    try:
        cantidad_int = int(cantidad)
        num_pagina_int = int(num_pagina)
        anio_int = int(anio_edicion[:4])
    except ValueError:
        return jsonify({'msg': 'Cantidad y paginas deben ser numericas', 'icono': 'warning'})
    if cantidad_int <= 0 or num_pagina_int <= 0:
        return jsonify({'msg': 'Cantidad y paginas deben ser mayores a 0', 'icono': 'warning'})
    if anio_int < 1000 or anio_int > current_year:
        return jsonify({'msg': 'El anio de edicion debe estar entre 1000 y el anio actual', 'icono': 'warning'})

    img = request.files.get('imagen')
    url_imagen = str_clean(request.form.get('url_imagen', ''))
    if url_imagen and not SecurityUtils.validate_url(url_imagen):
        return jsonify({'msg': 'La URL de la imagen no es valida', 'icono': 'warning'})
    img_nombre = 'logo.png'
    upload_folder = os.path.join(current_app.root_path, 'static', 'img', 'libros')
    os.makedirs(upload_folder, exist_ok=True)

    if img and img.filename:
        ext = img.filename.rsplit('.', 1)[-1].lower()
        is_image = img.content_type and img.content_type.startswith('image/')
        allowed_exts = ('png', 'jpg', 'jpeg', 'webp', 'gif', 'avif', 'svg', 'bmp', 'tiff', 'jfif', 'ico', 'heic', 'heif')
        if not is_image and ext not in allowed_exts:
            return jsonify({'msg': 'Archivo no permitido', 'icono': 'warning'})
        img_nombre = datetime.now().strftime('%Y%m%d%H%M%S') + f'.{ext}'
    elif url_imagen:
        import urllib.request
        from urllib.parse import urlparse
        try:
            parsed = urlparse(url_imagen)
            ext = 'jpg'
            if '.' in parsed.path:
                ext_guess = parsed.path.rsplit('.', 1)[-1].lower()
                is_image_guess = ext_guess in ('png', 'jpg', 'jpeg', 'webp', 'gif', 'avif', 'svg', 'bmp', 'tiff', 'jfif', 'ico', 'heic', 'heif')
                if is_image_guess:
                    ext = ext_guess
            img_nombre = datetime.now().strftime('%Y%m%d%H%M%S') + f'_url.{ext}'
            path_temp = os.path.join(upload_folder, img_nombre)
            req = urllib.request.Request(url_imagen, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=10) as response, open(path_temp, 'wb') as out_file:
                out_file.write(response.read())
        except Exception as e:
            return jsonify({'msg': 'No se pudo descargar la imagen de la URL', 'icono': 'warning'})
    elif request.form.get('foto_actual'):
        img_nombre = request.form.get('foto_actual')

    if id_ == '':
        data = m.insertar_libros(titulo, id_autor, id_editorial, id_materia, cantidad_int,
                                 num_pagina_int, anio_edicion, descripcion, img_nombre)
        if data == 'ok':
            if img and img.filename:
                img.save(os.path.join(upload_folder, img_nombre))
            return jsonify({'msg': 'Libro registrado', 'icono': 'success'})
        elif data == 'existe':
            return jsonify({'msg': 'El libro ya existe', 'icono': 'warning'})
        return jsonify({'msg': 'Error al registrar', 'icono': 'error'})
    else:
        data = m.actualizar_libros(titulo, id_autor, id_editorial, id_materia, cantidad_int,
                                   num_pagina_int, anio_edicion, descripcion, img_nombre, id_)
        if data == 'modificado':
            if img and img.filename:
                img.save(os.path.join(upload_folder, img_nombre))
            return jsonify({'msg': 'Libro modificado', 'icono': 'success'})
        return jsonify({'msg': 'Error al modificar', 'icono': 'error'})


@libros_bp.route('/editar/<int:id_>')
@permission_required('Libros')
def editar(id_):
    redir = login_required()
    if redir:
        return redir
    m = LibrosModel()
    data = m.edit_libros(id_)
    if data:
        autor_row = m.select("SELECT autor FROM autor WHERE id = %s", (data['id_autor'],))
        editorial_row = m.select("SELECT editorial FROM editorial WHERE id = %s", (data['id_editorial'],))
        materia_row = m.select("SELECT materia FROM materia WHERE id = %s", (data['id_materia'],))
        data['autor'] = autor_row['autor'] if autor_row else ''
        data['editorial'] = editorial_row['editorial'] if editorial_row else ''
        data['materia'] = materia_row['materia'] if materia_row else ''
    return jsonify(data)


@libros_bp.route('/eliminar/<int:id_>')
@permission_required('Libros')
def eliminar(id_):
    redir = login_required()
    if redir:
        return redir
    m = LibrosModel()
    data = m.estado_libros(0, id_)
    if data == 1:
        return jsonify({'msg': 'Libro dado de baja', 'icono': 'success'})
    return jsonify({'msg': 'Error al eliminar', 'icono': 'error'})


@libros_bp.route('/reingresar/<int:id_>')
@permission_required('Libros')
def reingresar(id_):
    redir = login_required()
    if redir:
        return redir
    m = LibrosModel()
    data = m.estado_libros(1, id_)
    if data == 1:
        return jsonify({'msg': 'Libro restaurado', 'icono': 'success'})
    return jsonify({'msg': 'Error al restaurar', 'icono': 'error'})


@libros_bp.route('/verificar/<int:id_>')
@permission_required('Libros')
def verificar(id_):
    redir = login_required()
    if redir:
        return redir
    m = LibrosModel()
    data = m.edit_libros(id_)
    if data:
        return jsonify({'cantidad': data['cantidad'], 'icono': 'success'})
    return jsonify({'msg': 'Error Fatal', 'icono': 'error'})


@libros_bp.route('/buscarLibro')
@permission_required('Libros')
def buscar_libro():
    redir = login_required()
    if redir:
        return redir
    valor = request.args.get('lb', '')
    m = LibrosModel()
    return jsonify(m.buscar_libro(valor))


@libros_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir: return redir
    from fpdf import FPDF
    from flask import Response
    id_autor     = request.args.get('id_autor')
    id_editorial = request.args.get('id_editorial')
    id_materia   = request.args.get('id_materia')
    estado       = request.args.get('estado')
    stock        = request.args.get('stock')
    from models.configuracion_model import ConfiguracionModel
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    m = LibrosModel()
    registros = m.get_libros(id_autor=id_autor, id_editorial=id_editorial,
                              id_materia=id_materia, estado=estado, stock=stock)
    # Construir etiqueta de filtros activos
    filtros_txt = []
    if id_autor:     filtros_txt.append(f'Autor ID:{id_autor}')
    if id_editorial: filtros_txt.append(f'Editorial ID:{id_editorial}')
    if id_materia:   filtros_txt.append(f'Materia ID:{id_materia}')
    if estado != '' and estado is not None:
        filtros_txt.append('Activos' if str(estado) == '1' else 'Inactivos')
    subtitulo = 'Filtros: ' + ', '.join(filtros_txt) if filtros_txt else 'Todos los registros'
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Listado de Libros')
    pdf_doc.set_font('Arial', 'B', 14)
    pdf_doc.cell(195, 8, datos.get('nombre', 'Biblioteca') + ' - Listado de Libros', 0, 1, 'C')
    pdf_doc.set_font('Arial', 'I', 9)
    pdf_doc.set_text_color(100, 100, 100)
    pdf_doc.cell(195, 5, subtitulo, 0, 1, 'C')
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.ln(3)
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(6, 11, 20)
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(10, 6, 'N°', 1, 0, 'C', True)
    pdf_doc.cell(55, 6, 'Titulo', 1, 0, 'C', True)
    pdf_doc.cell(40, 6, 'Autor', 1, 0, 'C', True)
    pdf_doc.cell(40, 6, 'Materia', 1, 0, 'C', True)
    pdf_doc.cell(15, 6, 'Cant.', 1, 0, 'C', True)
    pdf_doc.cell(20, 6, 'Estado', 1, 1, 'C', True)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(registros, 1):
        pdf_doc.cell(10, 6, str(i), 1, 0, 'C')
        pdf_doc.cell(55, 6, str(row.get('titulo', '')), 1, 0, 'L')
        pdf_doc.cell(40, 6, str(row.get('autor', '')), 1, 0, 'L')
        pdf_doc.cell(40, 6, str(row.get('materia', '')), 1, 0, 'L')
        pdf_doc.cell(15, 6, str(row.get('cantidad', '')), 1, 0, 'C')
        pdf_doc.cell(20, 6, 'Activo' if row.get('estado') == 1 else 'Inactivo', 1, 1, 'C')
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=reporte.pdf'})


@libros_bp.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response
    id_autor     = request.args.get('id_autor')
    id_editorial = request.args.get('id_editorial')
    id_materia   = request.args.get('id_materia')
    estado       = request.args.get('estado')
    stock        = request.args.get('stock')
    from models.configuracion_model import ConfiguracionModel
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    m = LibrosModel()
    registros = m.get_libros(id_autor=id_autor, id_editorial=id_editorial,
                              id_materia=id_materia, estado=estado, stock=stock)
    filtros_txt = []
    if id_autor:     filtros_txt.append(f'Autor ID:{id_autor}')
    if id_editorial: filtros_txt.append(f'Editorial ID:{id_editorial}')
    if id_materia:   filtros_txt.append(f'Materia ID:{id_materia}')
    if estado != '' and estado is not None:
        filtros_txt.append('Activos' if str(estado) == '1' else 'Inactivos')
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Listado de Libros"
    # Fila de filtros activos
    subtitulo = f"{datos.get('nombre', 'Biblioteca')} | Filtros: " + (', '.join(filtros_txt) if filtros_txt else 'Todos los registros')
    ws.merge_cells('A1:F1')
    sub_cell = ws['A1']
    sub_cell.value = subtitulo
    sub_cell.font = Font(italic=True, color='FF555555', size=10)
    sub_cell.alignment = Alignment(horizontal='center')
    header_fill = PatternFill("solid", fgColor="060b14")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers_arr = ['N°', 'Título', 'Autor', 'Materia', 'Cant.', 'Estado']
    for col, h in enumerate(headers_arr, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    for i, row in enumerate(registros, 1):
        ws.append([i, row.get('titulo', ''), row.get('autor', ''), row.get('materia', ''), row.get('cantidad', ''), 'Activo' if row.get('estado') == 1 else 'Inactivo'])
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.xlsx'}
    )


@libros_bp.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from flask import Response
    id_autor     = request.args.get('id_autor')
    id_editorial = request.args.get('id_editorial')
    id_materia   = request.args.get('id_materia')
    estado       = request.args.get('estado')
    stock        = request.args.get('stock')
    from models.configuracion_model import ConfiguracionModel
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    m = LibrosModel()
    registros = m.get_libros(id_autor=id_autor, id_editorial=id_editorial,
                              id_materia=id_materia, estado=estado, stock=stock)
    filtros_txt = []
    if id_autor:     filtros_txt.append(f'Autor ID:{id_autor}')
    if id_editorial: filtros_txt.append(f'Editorial ID:{id_editorial}')
    if id_materia:   filtros_txt.append(f'Materia ID:{id_materia}')
    if estado != '' and estado is not None:
        filtros_txt.append('Activos' if str(estado) == '1' else 'Inactivos')
    subtitulo = 'Filtros: ' + ', '.join(filtros_txt) if filtros_txt else 'Todos los registros'
    doc = Document()
    title = doc.add_heading(datos.get('nombre', 'Biblioteca') + ' - Listado de Libros', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Generado el {date.today().isoformat()} | {subtitulo}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    headers_arr = ['N°', 'Título', 'Autor', 'Materia', 'Cant.', 'Estado']
    table = doc.add_table(rows=1, cols=len(headers_arr))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers_arr):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs: run.bold = True
    for i, row in enumerate(registros, 1):
        cells = table.add_row().cells
        vals = [str(i), str(row.get('titulo', '')), str(row.get('autor', '')), str(row.get('materia', '')), str(row.get('cantidad', '')), 'Activo' if row.get('estado') == 1 else 'Inactivo']
        for j, v in enumerate(vals):
            cells[j].text = v
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    today = date.today().isoformat()
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=reporte_{today}.docx'}
    )


@libros_bp.route('/plantilla_excel')
@permission_required('Libros')
def plantilla_excel():
    """Descarga una plantilla Excel vacia con el formato requerido para carga masiva."""
    redir = login_required()
    if redir: return redir
    from io import BytesIO
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from flask import Response

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Carga_Masiva"

    headers_arr = ['Titulo', 'Autor', 'Editorial', 'Materia', 'Cantidad', 'Num Paginas', 'Anio Edicion', 'Descripcion']
    header_fill = PatternFill("solid", fgColor="060b14")
    header_font = Font(bold=True, color="FFFFFF", size=11)

    for col, h in enumerate(headers_arr, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')

    ejemplo = ['Cien Anios de Soledad', 'Gabriel Garcia Marquez', 'Sudamericana', 'Literatura', 3, 471, 1967, 'Novela clasica de la literatura latinoamericana']
    for col, val in enumerate(ejemplo, 1):
        ws.cell(row=2, column=col, value=val)

    anchos = [40, 30, 25, 20, 12, 14, 14, 40]
    for i, ancho in enumerate(anchos, 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = ancho

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename=plantilla_carga_masiva.xlsx'}
    )


@libros_bp.route('/upload_excel', methods=['POST'])
@permission_required('Libros')
def upload_excel():
    """Procesa un archivo Excel y realiza la carga masiva de libros."""
    redir = login_required()
    if redir: return redir
    import openpyxl
    from io import BytesIO

    archivo = request.files.get('archivo_excel')
    if not archivo or not archivo.filename.endswith('.xlsx'):
        return jsonify({'msg': 'Por favor sube un archivo Excel (.xlsx) valido.', 'icono': 'warning'})

    try:
        contenido = archivo.read()
        wb = openpyxl.load_workbook(BytesIO(contenido), data_only=True)
        ws = wb.active
    except Exception:
        return jsonify({'msg': 'No se pudo leer el archivo. Verifica que sea un Excel valido.', 'icono': 'error'})

    insertados = 0
    sumados = 0
    errores = 0
    detalles_errores = []

    m = LibrosModel()

    for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not any(row):
            continue

        try:
            titulo      = str(row[0]).strip() if row[0] else ''
            autor       = str(row[1]).strip() if row[1] else ''
            editorial   = str(row[2]).strip() if row[2] else ''
            materia     = str(row[3]).strip() if row[3] else ''
            cantidad    = int(row[4]) if row[4] else 0
            num_pagina  = int(row[5]) if row[5] else 0
            anio_raw    = str(int(row[6])) if isinstance(row[6], (float, int)) else (str(row[6]).strip() if row[6] else '')
            descripcion = str(row[7]).strip() if row[7] else ''

            if not titulo or not autor or not editorial or not materia:
                detalles_errores.append(f'Fila {idx}: Titulo, Autor, Editorial y Materia son obligatorios.')
                errores += 1
                continue
            if cantidad <= 0 or num_pagina <= 0:
                detalles_errores.append(f'Fila {idx}: Cantidad y Paginas deben ser mayores a 0.')
                errores += 1
                continue

            anio_edicion = normalize_anio_edicion(anio_raw)
            if not anio_edicion:
                detalles_errores.append(f'Fila {idx}: Anio invalido "{anio_raw}". Usa AAAA o AAAA-MM-DD.')
                errores += 1
                continue

            id_autor     = m.get_or_create_autor(autor)
            id_editorial = m.get_or_create_editorial(editorial)
            id_materia   = m.get_or_create_materia(materia)

            if not id_autor or not id_editorial or not id_materia:
                detalles_errores.append(f'Fila {idx}: No se pudieron crear las dependencias para "{titulo}".')
                errores += 1
                continue

            resultado = m.insertar_o_sumar_libro(titulo, id_autor, id_editorial, id_materia,
                                                  cantidad, num_pagina, anio_edicion, descripcion)
            if resultado == 'insertado':
                insertados += 1
            elif resultado == 'sumado':
                sumados += 1
            else:
                detalles_errores.append(f'Fila {idx}: Error al procesar "{titulo}".')
                errores += 1

        except Exception as e:
            detalles_errores.append(f'Fila {idx}: Error inesperado - {str(e)}')
            errores += 1

    partes = []
    if insertados: partes.append(f'{insertados} libro(s) insertado(s)')
    if sumados:    partes.append(f'{sumados} libro(s) con cantidad sumada')
    if errores:    partes.append(f'{errores} fila(s) con error')

    resumen = ', '.join(partes) if partes else 'No se proceso ninguna fila.'
    icono = 'success' if (insertados + sumados) > 0 and errores == 0 else \
            ('warning' if (insertados + sumados) > 0 and errores > 0 else 'error')

    return jsonify({
        'msg': resumen,
        'icono': icono,
        'detalles': detalles_errores
    })

@libros_bp.route('/run_image_updater')
@permission_required('Libros')
def run_image_updater():
    redir = login_required()
    if redir: return redir
    if session.get('id_usuario') != 1:
        return jsonify({'msg': 'Sin permisos', 'icono': 'error'})
    
    import subprocess
    import sys
    import os
    
    # Path to the actual script in the root folder
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    script_path = os.path.join(base_dir, 'actualizar_imagenes.py')
    
    python_exec = sys.executable or 'py'
    
    if os.name == 'nt':
        # Ejecutar en segundo plano sin mostrar la ventana CMD negra
        CREATE_NO_WINDOW = 0x08000000
        
        # Eliminar archivo de progreso anterior para evitar cruces
        progreso_file = os.path.join(base_dir, 'temp', 'progreso_imagenes.json')
        if os.path.exists(progreso_file):
            try: os.remove(progreso_file)
            except: pass
            
        subprocess.Popen([python_exec, script_path], creationflags=CREATE_NO_WINDOW, cwd=base_dir)
    else:
        progreso_file = os.path.join(base_dir, 'temp', 'progreso_imagenes.json')
        if os.path.exists(progreso_file):
            try: os.remove(progreso_file)
            except: pass
        subprocess.Popen([python_exec, script_path], cwd=base_dir)
        
    return jsonify({'msg': 'El escáner de imágenes se ha iniciado y se está procesando en segundo plano. Recargue la página en unos minutos para ver los cambios.', 'icono': 'success'})

@libros_bp.route('/progreso_imagenes')
@permission_required('Libros')
def progreso_imagenes():
    redir = login_required()
    if redir: return redir
    import os, json
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    progreso_file = os.path.join(base_dir, 'temp', 'progreso_imagenes.json')
    if os.path.exists(progreso_file):
        try:
            with open(progreso_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return jsonify(data)
        except Exception:
            return jsonify({'estado': 'error', 'mensaje': 'Leyendo estado...'})
    return jsonify({'estado': 'no_iniciado', 'mensaje': 'Esperando inicio...'})

@libros_bp.route('/api/chat', methods=['POST'])
def api_chat():
    if not session.get('activo'):
        return jsonify({'reply': 'Sesión no activa. Por favor inicia sesión de nuevo.', 'suggestions': []}), 401
        
    req_data = request.get_json() or {}
    message_raw = req_data.get('message', '').strip().lower()
    
    # Normalizar acentos
    import unicodedata
    message = "".join(c for c in unicodedata.normalize('NFD', message_raw) if unicodedata.category(c) != 'Mn')
    
    m = LibrosModel()
    
    # 1. Total books / cuántos libros
    if 'cuanto' in message and 'libro' in message:
        res_sum = m.select("SELECT COALESCE(SUM(cantidad), 0) as total_qty, COUNT(*) as total_titles FROM libro WHERE estado = 1")
        qty = res_sum['total_qty'] if res_sum else 0
        titles = res_sum['total_titles'] if res_sum else 0
        reply = f"En la biblioteca contamos con {titles} títulos de libros diferentes y un total de {qty} ejemplares en el catálogo."
        suggestions = ["Categorías", "¿Qué libros tienes?", "Préstamos activos"]
        
    # 2. Categories / Categorías / Materias
    elif 'categoria' in message or 'materia' in message:
        res_cat = m.select_all("SELECT m.materia, COUNT(l.id) as cant FROM materia m LEFT JOIN libro l ON m.id = l.id_materia WHERE m.estado = 1 GROUP BY m.materia ORDER BY cant DESC LIMIT 8")
        if res_cat:
            cats = [f"{r['materia']} ({r['cant']})" for r in res_cat]
            reply = "Las principales categorías disponibles y sus cantidades de libros son: " + ", ".join(cats) + "."
        else:
            reply = "No tenemos categorías registradas de momento."
        suggestions = ["¿Cuántos libros hay?", "¿Qué libros tienes?", "Préstamos activos"]
        
    # 3. Titles / Qué libros / Ver títulos
    elif 'que libro' in message or 'ver titulo' in message or 'lista' in message or 'titulo' in message:
        res_books = m.select_all("SELECT l.titulo, a.autor FROM libro l JOIN autor a ON l.id_autor = a.id WHERE l.estado = 1 ORDER BY l.id DESC LIMIT 5")
        if res_books:
            titles = [f"«{r['titulo']}» de {r['autor']}" for r in res_books]
            reply = "Aquí tienes algunos de nuestros libros disponibles recientemente: " + "; ".join(titles) + "."
        else:
            reply = "No hay libros disponibles de momento."
        suggestions = ["¿Cuántos libros hay?", "Categorías", "Préstamos activos"]
        
    # 4. Loans / Préstamos activos / Prestados
    elif 'prestamo' in message or 'prestado' in message or 'vencido' in message:
        res_loans = m.select("SELECT COUNT(*) as total FROM prestamo WHERE estado = 'Activo'")
        count = res_loans['total'] if res_loans else 0
        reply = f"Actualmente hay {count} préstamo(s) activo(s) en el sistema."
        suggestions = ["¿Cuántos libros hay?", "Categorías", "¿Qué libros tienes?"]
        
    # 5. Search book by general query/keyword
    else:
        res_search = m.select_all(
            "SELECT l.titulo, a.autor, l.cantidad FROM libro l "
            "JOIN autor a ON l.id_autor = a.id "
            "WHERE (l.titulo ILIKE %s OR a.autor ILIKE %s) AND l.estado = 1 LIMIT 5",
            (f"%{message_raw}%", f"%{message_raw}%")
        )
        if res_search:
            books_list = [f"«{r['titulo']}» de {r['autor']} ({r['cantidad']} disponibles)" for r in res_search]
            reply = f"Encontré estos libros relacionados: " + "; ".join(books_list) + "."
            suggestions = ["¿Cuántos libros hay?", "Categorías", "¿Qué libros tienes?"]
        else:
            reply = f"No logré encontrar ningún libro o autor que coincida con '{message_raw}'. Intenta buscar por otro nombre o pregúntame sobre el total de libros, categorías o préstamos activos."
            suggestions = ["¿Cuántos libros hay?", "Categorías", "¿Qué libros tienes?", "Préstamos activos"]
            
    return jsonify({
        'reply': reply,
        'suggestions': suggestions
    })